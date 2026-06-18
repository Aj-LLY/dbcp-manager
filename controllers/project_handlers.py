"""
项目事件处理函数模块 - 从 MainWindow 中抽取的项目 CRUD 相关事件处理器

本模块包含与项目卡片操作相关的所有事件处理函数，采用函数式风格：
  - 每个函数以 MainWindow 实例作为第一个参数（main_window）
  - 通过 main_window 访问所有服务层和 UI 组件（._project_service 等）

主要功能：
  1. 项目增删改查（新增、删除、编辑、复制）
  2. 卡片交互（选中、详情、阶段移动、列宽调整）
  3. 项目文件夹结构创建与 NDA 模板生成

设计原则：
  - 单一职责：每个函数只处理一种事件，逻辑清晰可测
  - 依赖注入：所有服务通过 main_window 注入，避免全局状态
  - 防御性编程：关键操作前检查前置条件（如阶段存在性、项目存在性）
"""

# =============================================================================
# 导入区
# =============================================================================

import os  # 操作系统接口，用于创建目录和处理路径
import shutil  # 文件操作工具，用于复制模板文件
import tkinter as tk  # Python 标准 GUI 库
from tkinter import messagebox  # 消息弹窗组件
from datetime import date  # 日期类，用于生成日期字符串

from models.project import Project  # 项目实体类（用于类型标注）
from ui.project_card import ProjectCard  # 项目卡片组件（用于类型标注）
from ui.project_dialog import show_project_dialog  # 项目编辑对话框便捷函数
from ui.detail_dialog import show_detail_dialog  # 项目详情对话框便捷函数
from utils.config import Config  # 全局配置类


# =============================================================================
# 工具栏事件处理函数 - 响应顶部工具栏各按钮的点击
# =============================================================================

def on_add_project(main_window):
    """处理"新增项目"按钮点击事件。

    执行流程：
      1. 检查是否有已配置的流程阶段（无阶段时无法创建项目，提示用户先配置）
      2. 打开新增项目对话框，预填当前所有流程阶段供下拉选择
      3. 用户填写并确认后，通过 ProjectService 创建项目实体
      4. 创建成功后自动建立项目文件夹结构，刷新看板显示

    文件夹路径处理逻辑：
      - 如果用户在对话框中指定了已存在的文件夹路径，直接使用该路径
      - 否则调用 create_project_folder() 自动创建标准目录结构

    数据流：
      用户输入 -> show_project_dialog() 返回 dict -> ProjectService.create_project() -> 写入 JSON

    Args:
        main_window: MainWindow 实例，提供以下访问入口：
            - _workflow_service: 流程阶段服务（获取阶段列表）
            - _project_service: 项目业务服务（创建项目）
            - _data_service: 数据持久化服务（更新项目字段）
            - _refresh_kanban: 看板刷新方法

    Returns:
        None: 该函数通过副作用（弹窗、数据持久化、UI刷新）完成操作，无返回值

    Raises:
        无显式抛出异常：所有异常在内部通过 messagebox 处理并向用户展示
    """
    # 前置条件检查：获取所有已配置的流程阶段
    stages = main_window._workflow_service.get_all_stages()  # 从流程服务获取所有已配置的阶段列表
    if not stages:  # 如果没有任何流程阶段
        messagebox.showwarning("提示", "请先配置流程阶段")  # 弹窗提示用户需要先配置阶段
        return  # 中止操作，不打开对话框

    # 打开新增项目对话框，传入阶段列表供用户选择初始阶段
    result = show_project_dialog(main_window, "新增项目", stages=stages)  # 返回 dict 或 None（取消时）
    if result:  # 用户点击了"确认"按钮（非取消），result 为包含所有表单字段的字典
        # 调用项目服务创建项目，传入用户在对话框中填写的各字段值
        success, msg, project = main_window._project_service.create_project(
            company_name=result["company_name"],  # 客户公司名称（必填）
            system_name=result["system_name"],  # 被测系统名称（必填）
            cert_number=result["cert_number"],  # 证书备案编号（可选）
            issue_date=result.get("issue_date", ""),  # 证书签发日期（可选）
            level=result.get("level", ""),  # 系统安全保护等级（可选）
            location=result.get("location", ""),  # 项目所属地（可选）
            deadline=result["deadline"],  # 项目截止交付日期（必填）
            notes=result["notes"],  # 备注信息（可选）
            stage_id=result["stage_id"],  # 初始流程阶段 ID（必填）
        )
        if success:
            # 创建项目文件夹结构（或使用用户指定的已有目录）
            user_folder = result.get("folder_path", "").strip()
            if user_folder and os.path.isdir(user_folder):
                # 用户指定了已有目录，直接使用该路径
                project.folder_path = user_folder
                main_window._data_service.update_project(project.id, {"folder_path": user_folder})
            else:
                # 自动创建标准项目文件夹结构（含子目录和 NDA 模板）
                create_project_folder(main_window, project)
            main_window._refresh_kanban()  # 刷新看板，显示新创建的项目卡片
        else:
            messagebox.showerror("错误", msg)  # 创建失败时弹窗显示错误原因


def on_delete_selected(main_window):
    """处理"删除项目"按钮点击事件。

    执行流程：
      1. 获取当前看板中选中的项目 ID
      2. 如果未选中任何卡片，提示用户先选择项目
      3. 显示二次确认对话框（防止误删，不可撤销）
      4. 确认后通过 ProjectService 执行永久删除
      5. 删除成功后刷新看板

    安全设计：
      - 二次确认机制：显示项目名称并要求用户明确确认
      - 明确标注"此操作不可撤销"以提醒用户谨慎操作
      - 防御性检查：删除前验证项目是否仍然存在

    Args:
        main_window: MainWindow 实例，提供以下访问入口：
            - _kanban: 看板 UI 组件（获取选中项目 ID）
            - _project_service: 项目业务服务（查询和删除项目）
            - _refresh_kanban: 看板刷新方法

    Returns:
        None

    Raises:
        无显式抛出异常
    """
    # 从看板获取当前高亮选中卡片的项目 ID
    project_id = main_window._kanban.get_selected_project_id()
    if not project_id:  # 没有任何卡片被选中
        messagebox.showinfo("提示", "请先在卡片上点击选择要删除的项目")  # 提示用户先选择
        return  # 中止操作

    # 根据 ID 获取项目对象，用于在确认对话框中显示项目名称
    project = main_window._project_service.get_project_by_id(project_id)
    if not project:  # 防御性检查：项目可能已被其他操作删除
        return

    # 二次确认对话框：显示项目名称，提醒用户此操作不可撤销
    # \u300c = 「，\u300d = 」—— 中文书名号式的引号包裹项目名
    if messagebox.askyesno("确认删除",
                           f"确定要永久删除项目\u300c{project.name}\u300d吗？\n\n"
                           "此操作不可撤销！"):
        # 通过项目服务执行永久删除（同时清理关联数据）
        success, msg = main_window._project_service.delete_project(project_id)
        if success:  # 删除成功
            main_window._refresh_kanban()  # 刷新看板，移除已删除的项目卡片
        else:
            messagebox.showerror("错误", msg)  # 删除失败时弹窗显示错误信息


# =============================================================================
# 看板事件处理函数 - 响应看板中卡片的交互操作
# =============================================================================

def on_card_selected(main_window, card: ProjectCard | None):
    """处理卡片选中状态变化事件。

    当用户单击看板中的卡片时触发。可用于扩展功能，如：
      - 更新状态栏显示当前选中项目的信息
      - 启用/禁用依赖于选中状态的操作按钮（如删除按钮）

    当前为空实现，预留扩展点。遵循 YAGNI 原则，仅在需要时添加逻辑。

    Args:
        main_window: MainWindow 实例
        card: 当前选中的卡片组件实例，或 None（表示取消选中/点击空白区域）
            - ProjectCard: 用户点击了某个卡片，该卡片进入选中状态
            - None: 用户点击了看板空白区域，取消所有选中

    Returns:
        None
    """
    pass  # 预留扩展：可在此添加状态栏更新、右键菜单等逻辑


def pick_project_from_card(main_window, card: ProjectCard) -> Project | None:
    """从合并卡片中选择单个项目。

    当一张卡片合并显示了多个系统时（同一公司多个系统共享卡片），
    需要让用户选择具体要操作哪个系统的项目。

    行为逻辑：
      - 单项目卡片（card.projects 为空或仅含 1 项）：直接返回该项目，无需弹窗
      - 多项目卡片（card.projects 包含 2 个以上项目）：弹出 Listbox 选择窗口

    Args:
        main_window: MainWindow 实例，用作弹出对话框的父窗口
        card: 被操作的合并卡片组件，包含一个或多个项目
            - card.projects: 合并卡片中的所有项目列表（多系统时存在）
            - card.project: 单项目时的项目引用

    Returns:
        Project | None:
            - Project: 用户选中的项目实体（单选或从列表中选择）
            - None: 用户取消了选择操作（关闭对话框或未选中任何项）

    Raises:
        无显式抛出异常
    """
    # 获取卡片中的所有项目：优先取 projects 列表，降级取单个 project
    projects = card.projects or [card.project]
    if len(projects) == 1:
        # 单项目卡片：无需弹窗，直接返回
        return projects[0]

    # --- 多项目卡片：构建选择对话框 ---
    dlg = tk.Toplevel(main_window)  # 创建模态子窗口
    dlg.title("选择系统")
    dlg.geometry("320x250")
    dlg.configure(bg="#ffffff")  # 白色背景
    dlg.grab_set()  # 设为模态，阻塞对主窗口的操作

    # 标题提示文字
    tk.Label(dlg, text="请选择要操作的系统:", bg="#ffffff",
             font=("Microsoft YaHei", 11, "bold")).pack(pady=(15, 10))

    # 可滚动列表：列出所有系统名称
    lb = tk.Listbox(dlg, font=("Microsoft YaHei", 10), selectmode="single")  # 单选模式
    lb.pack(fill=tk.BOTH, expand=True, padx=20, pady=5)
    for p in projects:
        # 列表项显示系统名称（优先用 system_name，降级用 name）
        lb.insert(tk.END, p.system_name or p.name)
    lb.selection_set(0)  # 默认选中第一项

    # 使用可变容器在闭包中传递选中结果
    result = {"project": None}

    def _ok():
        """确定按钮回调：将选中的项目写入 result 并关闭窗口"""
        sel = lb.curselection()  # 获取选中项的索引元组
        if sel:
            result["project"] = projects[sel[0]]  # 取第一个选中索引对应的项目
        dlg.destroy()  # 关闭窗口

    # --- 底部按钮区域 ---
    btn_frame = tk.Frame(dlg, bg="#f0f2f5")
    btn_frame.pack(fill=tk.X, side=tk.BOTTOM)
    tk.Frame(btn_frame, bg="#d0d5dd", height=1).pack(fill=tk.X)  # 顶部分割线
    inner = tk.Frame(btn_frame, bg="#f0f2f5")
    inner.pack(fill=tk.X, padx=16, pady=8)

    # 取消按钮（白色背景，靠右）
    tk.Button(inner, text="取消", command=dlg.destroy,
        bg="#ffffff", fg="#2c3e50", cursor="hand2").pack(side=tk.RIGHT, padx=(10,0))

    # 确定按钮（蓝色背景，白色文字，靠右）
    tk.Button(inner, text="确定", command=_ok,
        bg="#3498db", fg="white", cursor="hand2").pack(side=tk.RIGHT)

    # 绑定回车键快捷确认
    dlg.bind("<Return>", lambda e: _ok())

    # 阻塞等待对话框关闭（wait_window 会挂起当前执行流直到 dlg 销毁）
    main_window.wait_window(dlg)
    return result["project"]


def on_card_detail(main_window, card: ProjectCard):
    """处理卡片"详情"按钮点击事件。

    打开项目详情对话框，展示项目的完整信息（阶段历史、操作日志等）。
    同时支持在详情对话框中编辑和删除项目。

    详情对话框支持的交互：
      - 查看：展示项目所有字段、阶段变更历史、操作日志
      - 编辑：修改项目字段并保存
      - 删除：永久删除项目
      - 阶段移动：在详情对话框中直接切换流程阶段
      - 多系统：如果卡片合并了多个系统，同时展示所有系统信息

    Args:
        main_window: MainWindow 实例，提供以下访问入口：
            - _workflow_service: 流程服务（获取阶段列表）
            - _log_service: 日志服务（获取操作日志）
            - _project_service: 项目服务（查询、更新、删除、移动项目）
            - _refresh_kanban: 看板刷新方法
        card: 被点击"详情"按钮的项目卡片组件
            - card.projects: 合并卡片中的项目列表（多系统时）
            - card.project: 单项目卡片中的项目引用

    Returns:
        None
    """
    # 获取项目引用：多系统卡片取第一个，单系统卡片直接取
    project = card.projects[0] if card.projects else card.project
    # 如果有多个系统，传递完整列表供详情对话框展示
    all_projects = card.projects if card.projects and len(card.projects) > 1 else None

    # 预加载阶段列表和操作日志，传递给详情对话框
    stages = main_window._workflow_service.get_all_stages()
    logs = main_window._log_service.get_project_logs(project.id)

    def _handle_move(target_stage_id, dialog):
        """移动阶段回调：在详情对话框中点击箭头按钮时触发。

        Args:
            target_stage_id: 目标阶段 ID
            dialog: 详情对话框实例（用于刷新数据）
        """
        # 通过项目服务执行阶段移动
        success, msg = main_window._project_service.move_project(project.id, target_stage_id)
        if success:
            main_window._refresh_kanban()  # 刷新看板以反映阶段变更
            # 获取更新后的项目数据并刷新详情对话框
            upd = main_window._project_service.get_project_by_id(project.id)
            dialog.refresh_data(upd,
                main_window._workflow_service.get_all_stages(),
                main_window._log_service.get_project_logs(project.id))
        else:
            messagebox.showerror("错误", msg)

    # 显示详情对话框，传入移动阶段回调
    result = show_detail_dialog(main_window, project, stages, logs,
                               on_move=_handle_move, all_projects=all_projects)
    if not result:
        # 用户关闭了对话框（未进行编辑或删除操作）
        return

    # 解析对话框返回的操作类型和数据
    action, data = result
    if action == "edit":
        systems = data.get("systems", [])
        proj_list = card.projects if card.projects and len(card.projects) > 1 else [project]
        print(f"[详情编辑] 原始项目={len(proj_list)} 返回系统={len(systems)}", flush=True)
        for i, s in enumerate(systems):
            print(f"[详情编辑]   sys#{i}: name={s.get('system_name')}", flush=True)
        for i, p in enumerate(proj_list):
            if i < len(systems):
                sys_data = systems[i]
                success, msg = main_window._project_service.update_project(
                    p.id,
                    company_name=data.get("company_name"),
                    system_name=sys_data.get("system_name", p.system_name),
                    cert_number=sys_data.get("cert_number", p.cert_number),
                    issue_date=sys_data.get("issue_date", p.issue_date),
                    level=sys_data.get("level", p.level),
                    location=data.get("location"),
                    deadline=data.get("deadline"),
                    notes=data.get("notes"),
                    stage_id=data.get("stage_id"),
                    folder_path=data.get("folder_path"),
                )
                if not success:
                    messagebox.showerror("错误", msg)
                    break
            else:
                print(f"[详情编辑] 删除多余项目: {p.system_name}", flush=True)
                main_window._project_service.delete_project(p.id)
        main_window._refresh_kanban()
    elif action == "delete":
        # 删除操作：永久删除项目
        success, msg = main_window._project_service.delete_project(project.id)
        if success:
            main_window._refresh_kanban()
        else:
            messagebox.showerror("错误", msg)


def on_card_edit(main_window, card: ProjectCard):
    """处理卡片"编辑"按钮点击事件（双击卡片也会触发此函数）。

    打开编辑项目对话框，允许用户修改项目的各字段值。
    支持单系统和多系统（合并卡片）两种模式：
      - 单系统：编辑单个项目的所有字段
      - 多系统：编辑共享字段（公司名、截止日期、备注、阶段、文件夹路径），
        同时为每个系统单独编辑系统名、证书编号、签发日期、等级、属地

    Args:
        main_window: MainWindow 实例，提供以下访问入口：
            - _workflow_service: 流程服务（获取阶段列表）
            - _project_service: 项目服务（更新项目）
            - _refresh_kanban: 看板刷新方法
        card: 被点击"编辑"按钮的项目卡片组件
            - card.projects: 合并卡片中的项目列表（多系统时存在）
            - card.project: 单项目卡片中的项目引用

    Returns:
        None
    """
    # 获取项目引用和完整项目列表
    project = card.projects[0] if card.projects else card.project
    all_proj = card.projects if card.projects and len(card.projects) > 1 else None

    # 获取阶段列表供编辑对话框使用
    stages = main_window._workflow_service.get_all_stages()

    # 显示编辑对话框，传入现有项目数据和项目列表
    result = show_project_dialog(main_window, "编辑项目", project, stages, all_projects=all_proj)
    if result:
        shared = {k: result.get(k) for k in ("company_name", "location",
            "deadline", "notes", "stage_id", "folder_path")}
        systems = result.get("systems", [])
        proj_list = card.projects if card.projects else [card.project]
        print(f"[卡片编辑] 原始项目={len(proj_list)} 返回系统={len(systems)}", flush=True)
        for i, s in enumerate(systems):
            print(f"[卡片编辑]   sys#{i}: name={s.get('system_name')}", flush=True)

        for i, p in enumerate(proj_list):
            if i < len(systems):
                sys_data = systems[i]
                success, msg = main_window._project_service.update_project(
                    p.id,
                    company_name=shared["company_name"],
                    system_name=sys_data.get("system_name", p.system_name),
                    cert_number=sys_data.get("cert_number", p.cert_number),
                    issue_date=sys_data.get("issue_date", p.issue_date),
                    level=sys_data.get("level", p.level),
                    location=shared["location"],
                    deadline=shared["deadline"],
                    notes=shared["notes"],
                    stage_id=shared["stage_id"],
                    folder_path=shared["folder_path"],
                )
            else:
                # 对话框返回的系统数少于原始项目数 → 删除多余项目
                print(f"[卡片编辑] 删除多余项目: {p.system_name}", flush=True)
                main_window._project_service.delete_project(p.id)

        main_window._refresh_kanban()


def on_card_move_stage(main_window, card: ProjectCard, target_stage_id: str):
    """处理卡片左箭头/右箭头的阶段移动事件。

    当用户点击卡片上的 ◀（左箭头）或 ▶（右箭头）时触发。
    左箭头将项目移到上一阶段，右箭头将项目移到下一阶段。

    性能优化：
      - 先在 UI 层面移动卡片（局部刷新），无需全量重建看板
      - 如果卡片合并了多个系统，同时移动所有系统

    防御性检查：
      - 如果源阶段和目标阶段相同（用户在边界处误点），直接返回，不执行任何操作

    Args:
        main_window: MainWindow 实例，提供以下访问入口：
            - _project_service: 项目服务（执行阶段移动）
            - _kanban: 看板 UI 组件（执行卡片移动动画）
        card: 被操作的项目卡片组件
            - card.projects: 合并卡片中的项目列表（多系统时）
            - card.project: 单项目卡片中的项目引用
        target_stage_id: 目标阶段的唯一标识符（由 KanbanBoard 根据箭头方向计算得出）

    Returns:
        None
    """
    # 防御性检查：源阶段与目标阶段相同则无需移动
    source_stage_id = card.project.stage_id
    if source_stage_id == target_stage_id:
        return

    # 移动合并卡片中的所有项目（单项目卡片退化为仅移动自身）
    projects = card.projects or [card.project]
    for p in projects:
        # 通过项目服务执行阶段移动（同时记录操作日志）
        main_window._project_service.move_project(p.id, target_stage_id)

    # 在 UI 层面移动卡片到目标列（性能优化：无需全量刷新）
    main_window._kanban.move_card_to_column(card, target_stage_id)


def on_card_copy(main_window, card: ProjectCard):
    """处理卡片"复制"按钮点击事件 - 创建当前项目的完整副本。

    复制逻辑：
      - 读取源项目的所有业务字段（公司名称、系统名称、编号、等级等）
      - 公司名称在原值后追加 " - 副本" 后缀以区分原项目
      - 通过 ProjectService.create_project() 创建新项目实体
      - 新项目与源项目处于同一流程阶段

    多系统支持：
      - 如果卡片合并了多个系统，每个系统都会创建一个独立的副本

    Args:
        main_window: MainWindow 实例，提供以下访问入口：
            - _project_service: 项目服务（创建项目副本）
            - _refresh_kanban: 看板刷新方法
        card: 被点击"复制"按钮的项目卡片组件
            - card.projects: 合并卡片中的项目列表（多系统时）
            - card.project: 单项目卡片中的项目引用

    Returns:
        None
    """
    # 遍历卡片中的所有项目（单项目退化为仅迭代自身）
    for p in (card.projects or [card.project]):
        # 在公司名称后追加 " - 副本" 后缀，\u526f\u672c = 副本
        copy_name = f"{p.company_name} - \u526f\u672c"
        # 创建新项目实体，复制源项目的所有业务字段，保持同一阶段
        main_window._project_service.create_project(
            company_name=copy_name,  # 带后缀的公司名
            system_name=p.system_name,  # 系统名保持不变
            cert_number=p.cert_number,  # 证书编号保持不变
            issue_date=p.issue_date,  # 签发日期保持不变
            level=p.level,  # 等级保持不变
            location=p.location,  # 属地保持不变
            deadline=p.deadline,  # 截止日期保持不变
            notes=p.notes,  # 备注保持不变
            stage_id=p.stage_id,  # 阶段 ID 保持不变
        )
    # 全量刷新看板以显示新创建的项目卡片
    main_window._refresh_kanban()


def on_column_resize(main_window, stage_id: str, new_width: int):
    """处理列宽拖拽完成事件 - 保存调整后的列宽到持久化数据。

    当用户拖拽看板列右侧的分隔手柄调整列宽后触发。
    将新宽度通过 WorkflowService 持久化保存，下次启动时自动恢复。

    职责分离：
      - UI 层（KanbanBoard）负责检测拖拽并计算新宽度
      - 本函数负责将宽度值持久化，不参与 UI 渲染逻辑

    Args:
        main_window: MainWindow 实例，提供 _workflow_service 访问入口
        stage_id: 被调整宽度列对应的阶段 ID
        new_width: 新的列宽值（单位：像素），由 KanbanBoard 计算

    Returns:
        None
    """
    # 将新列宽保存到持久化数据（JSON 文件），下次启动自动恢复
    main_window._workflow_service.update_stage_width(stage_id, new_width)


# =============================================================================
# 项目文件夹结构创建函数
# =============================================================================

def create_project_folder(main_window, project):
    """为新建项目创建完整的文件目录结构。

    在程序数据目录下创建以公司名+系统名+日期命名的文件夹，
    并初始化以下子目录：
      - 01-其他归档文件/                     # 存放杂项归档材料
      - 00-{公司}-{系统}-报告打印/            # 存放打印版报告（通过 generate_nda_template 创建）
      - 13-{公司}-{系统}-渗透测试报告/         # 存放渗透测试相关文件

    文件夹命名规则：
      格式：{序号}-{公司简称}-{系统简称}-{年月日}
      序号：当前项目总数（自动递增，格式化 3 位）
      日期：创建当日（格式 YYMMDD，如 260612）

    安全性处理：
      - 公司名称和系统名称中的路径分隔符（/ 和 \\）被替换为下划线
      - 使用 os.makedirs(exist_ok=True) 避免并发创建时抛异常

    创建成功后，将文件夹路径保存到 project 对象并持久化到 JSON 文件。

    Args:
        main_window: MainWindow 实例，提供以下访问入口：
            - _project_service: 项目服务（获取项目总数用于序号计算）
            - _data_service: 数据持久化服务（保存文件夹路径）
        project: 新创建的项目实体对象（包含公司名称、系统名称等信息）
            - project.company_name: 公司名称（用于目录名）
            - project.system_name: 系统名称（用于目录名）
            - project.id: 项目 ID（用于持久化文件夹路径关联）

    Returns:
        None

    Raises:
        无显式抛出异常：OSError 在内部静默捕获，避免因文件系统错误中断用户操作
    """
    try:
        # 获取程序数据根目录路径
        base = Config.get_data_dir()
        # 当前项目总数（含刚创建的），用于生成递增序号
        count = len(main_window._project_service.get_all_projects())
        # 当日日期，格式 YYMMDD（如 260612）
        date_str = date.today().strftime("%y%m%d")

        # 清理公司名称中的非法字符（路径分隔符会被误解析为子目录）
        cname = (project.company_name or "未命名").replace("/", "_").replace("\\", "_")
        # 清理系统名称中的非法字符
        sname = (project.system_name or "").replace("/", "_").replace("\\", "_")

        # 组装文件夹名称：序号(3位)-公司名-系统名-日期
        folder_name = f"{count:03d}-{cname}-{sname}-{date_str}"
        # 拼接完整文件夹路径
        root = os.path.join(base, folder_name)
        # 递归创建文件夹目录（exist_ok=True 避免并发创建时报 FileExistsError）
        os.makedirs(root, exist_ok=True)

        # 创建归档子目录结构
        archive_subdirs = [
            "01-其他归档文件/00-网安报备",
            "01-其他归档文件/01-备案材料",
            "01-其他归档文件/02-往期测评报告",
            "01-其他归档文件/03-现场测评",
            "01-其他归档文件/04-渗透漏扫",
        ]
        for d in archive_subdirs:
            os.makedirs(os.path.join(root, d), exist_ok=True)

        # 生成保密承诺书模板（复制模板文件并替换公司名称和日期）
        generate_nda_template(main_window, root, cname, project.company_name or "未命名")

        # 将文件夹路径关联到项目并持久化到 JSON
        project.folder_path = root  # 设置项目对象的文件夹路径属性
        main_window._data_service.update_project(project.id, {"folder_path": root})  # 持久化到 JSON
    except OSError:
        # 静默捕获文件系统错误（如磁盘满、权限不足），避免中断用户操作流程
        pass


def generate_nda_template(main_window, root, cname_clean, company_name):
    """生成保密承诺书 Word 模板，替换公司名称和日期占位符。

    从模板目录复制保密承诺书 .docx 文件到项目文件夹，
    并为当前项目自定义公司名称和签署日期。

    替换策略（处理 Word 文档中 run 级别的内容）：
      1. 逐 run 替换 "XX公司"/"xx公司" -> 实际公司名称
      2. 处理跨越 run 边界的 "XX" + "公司" 分割情况
      3. 在表格中替换日期占位符 "XX" -> 实际年月日
      4. 所有替换保持原始格式（字体、字号、颜色等样式不变）

    幂等性保证：
      - 目标文件已存在时跳过生成，避免覆盖已有文件
      - 模板文件不存在时静默跳过，不报错

    Args:
        main_window: MainWindow 实例（本函数中未直接使用，保留以保持接口一致性）
        root: 项目文件夹根路径，模板文件将被复制到此目录
        cname_clean: 清理后的公司名称（用于生成文件名，不含路径分隔符）
        company_name: 原始公司名称（用于替换文档内容中的占位符）

    Returns:
        None

    Raises:
        无显式抛出异常：所有异常在内部静默捕获，避免因模板处理失败中断项目创建流程
    """
    try:
        # 构建模板文件路径：data_dir/templates/02-保密承诺书模板.docx
        template_path = os.path.join(Config.get_data_dir(), "templates", "02-保密承诺书模板.docx")
        if not os.path.exists(template_path):
            # 模板文件不存在，静默跳过（不报错，不阻塞项目创建流程）
            return

        # 生成目标文件名：格式为 "02-{公司简称}-{公司全称}-保密承诺书.docx"
        dest_name = f"02-{cname_clean}-{company_name}-保密承诺书.docx" if company_name else f"02-{cname_clean}-保密承诺书.docx"
        dest_path = os.path.join(root, dest_name)

        if os.path.exists(dest_path):
            # 目标文件已存在，跳过生成（幂等性保护）
            return

        # 复制模板文件到项目文件夹（保留原文件的创建时间和权限信息）
        shutil.copy2(template_path, dest_path)

        # 导入 docx 库（延迟导入，避免全局依赖）
        import docx
        from datetime import datetime

        # 打开复制后的文件进行内容替换
        doc = docx.Document(dest_path)

        # --- 第一步：替换公司名称（段落级别） ---
        # 遍历所有段落，在 run 级别替换公司名称占位符
        for p in doc.paragraphs:
            for run in p.runs:
                # 替换中文 "XX公司" 和英文 "xx公司"
                if "XX公司" in run.text or "xx公司" in run.text:
                    run.text = run.text.replace("XX公司", company_name).replace("xx公司", company_name)
                    break  # 每个段落只替换第一个匹配的 run

        # --- 第二步：处理跨 run 边界的 "XX"+"公司" 分割情况 ---
        # 某些 Word 文档中 "XX公司" 被拆分为两个 run："XX" 和 "公司"
        for p in doc.paragraphs:
            for j in range(len(p.runs) - 1):
                if p.runs[j].text.strip() == "XX" and p.runs[j+1].text.strip() == "公司":
                    # 将公司名写入第一个 run，清空第二个 run
                    p.runs[j].text = company_name
                    p.runs[j+1].text = ""

        # --- 第三步：替换日期（表格级别） ---
        # 日期格式为 "XX年XX月XX日"，需要替换三个 "XX" 为实际年月日
        now = datetime.now()
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        rs = p.runs
                        # 检测格式：6 个 run，第 0/2/4 位为 "XX"
                        if len(rs) >= 6 and all(rs[k].text.strip() == v for k, v in [(0,"XX"),(2,"XX"),(4,"XX")]):
                            rs[0].text = str(now.year)  # 替换年份
                            rs[2].text = f"{now.month:02d}"  # 替换月份（补零）
                            rs[4].text = f"{now.day:02d}"  # 替换日期（补零）
                            break  # 每个单元格只替换第一处匹配

        # 保存修改后的文档
        doc.save(dest_path)
    except Exception:
        # 静默捕获所有异常：模板处理失败不应中断项目创建流程
        pass
