"""
项目卡片文件操作模块 - 从 project_card.py 提取的独立文件操作函数

本模块包含与项目文件夹交互的独立函数，原为 ProjectCard 类的方法。
提取为独立函数后，以 project 和 parent 参数替代 self 引用。

相关模块：
  - models.Project：项目数据实体类
  - utils.Config：全局 UI 配置（字体、颜色、尺寸等）
  - ui.dialog_report_print：报告打印对话框和 XLSX 生成

函数列表：
  - find_project_folder(project) -> str：查找项目文件夹路径
  - on_folder_click(project)：在文件管理器中打开项目文件夹
  - on_init_click(project, parent)：项目初始化（创建子目录和模板文件）
  - on_rename_click(project, parent)：批量重命名过程文件
  - on_zip_click(project, parent)：打包过程文档为 ZIP
  - on_report_print_click(project, parent)：报告打印按钮处理
"""

# =============================================================================
# 导入区
# =============================================================================

import os          # 文件系统操作
import re          # 正则匹配（文件名模式匹配）
import zipfile     # ZIP 压缩/解压
import shutil      # 高级文件操作（复制保留元数据）
import subprocess  # 子进程（打开文件管理器）
import sys         # 系统平台判断
from tkinter import messagebox  # 消息弹窗
from datetime import date       # 当前日期

# ---- 模型层 ----
from models.project import Project  # 项目实体类

# ---- 工具层 ----
from utils.config import Config  # 全局配置类

# ---- 报告打印对话框 ----
from ui.dialog_report_print import (
    show_report_dialog,         # 报告打印前的编辑确认对话框（14 个可编辑字段）
    _create_report_xlsx,        # 创建 XLSX 基础结构
    _create_report_xlsx_data,   # 根据编辑框数据创建 XLSX
)


# =============================================================================
# find_project_folder - 查找项目文件夹路径
# =============================================================================

def find_project_folder(project: Project) -> str:
    """根据项目信息查找本地文件夹路径

    查找策略（按优先级从高到低）：
      1. 项目存储的 folder_path 属性（最直接，优先使用）
      2. 按公司名 + 系统名 + 创建日期的关键词模糊搜索（兜底方案）

    搜索关键词：
      - 公司名称（清理路径非法字符后）
      - 系统名称（清理路径非法字符后）
      - 创建日期（YYMMDD 格式，取项目 created_at 前 10 位的后 6 位）

    Args:
        project: 项目实体对象

    Returns:
        str: 找到的文件夹路径，未找到返回空字符串 ""
    """
    # 策略 1：优先使用项目存储的文件夹路径
    if project.folder_path and os.path.isdir(project.folder_path):
        return project.folder_path  # 直接返回存储的路径

    # 策略 2：按关键词搜索
    base = Config.get_data_dir()  # 获取程序数据根目录
    if not os.path.exists(base):  # 根目录不存在
        return ""  # 无数据目录，返回空
    # 清理名称中的路径非法字符，统一替换为下划线
    cname = (project.company_name or "未命名").replace("/", "_").replace("\\", "_")
    sname = (project.system_name or "").replace("/", "_").replace("\\", "_")
    date_str = ""  # 日期关键词（默认为空，不参与过滤）
    if project.created_at:  # 有创建时间
        date_str = project.created_at[:10].replace("-", "")[2:]  # YYYY-MM-DD -> YYMMDD

    # 遍历数据目录，查找匹配的文件夹名
    for name in os.listdir(base):
        full = os.path.join(base, name)  # 拼接完整路径
        if not os.path.isdir(full):  # 非目录跳过
            continue
        # 匹配规则：名称中包含公司名，且（无系统名 或 包含系统名），且（无日期 或 包含日期）
        if cname in name and (not sname or sname in name) and (not date_str or date_str in name):
            return full  # 找到匹配，返回路径
    return ""  # 未找到任何匹配


# =============================================================================
# on_folder_click - 打开项目文件夹
# =============================================================================

def on_folder_click(project: Project):
    """处理"打开文件夹"按钮点击 - 在系统文件管理器中打开项目目录

    调用各操作系统的默认文件管理器：
      - Windows：os.startfile(path) 直接打开
      - Linux/macOS：subprocess.run(["xdg-open", path])

    Args:
        project: 项目实体对象
    """
    try:
        path = find_project_folder(project)  # 查找项目文件夹路径
        if path and os.path.isdir(path):  # 路径存在且是目录
            if sys.platform == "win32":  # Windows 系统
                os.startfile(path)  # 使用 Windows 默认方式打开（类似双击文件夹）
            else:  # 非 Windows 系统（Linux/macOS）
                subprocess.run(["xdg-open", path])  # 调用 xdg-open 打开
    except Exception:  # 打开失败（权限不足、路径不存在等）
        pass  # 静默处理：打开失败不阻塞主流程


# =============================================================================
# on_init_click - 项目初始化
# =============================================================================

def on_init_click(project: Project, parent=None):
    """项目初始化：创建标准子目录和保密承诺书模板

    创建内容：
      1. "01-其他归档文件" 子目录
      2. "00-{公司}-{系统}-报告打印" 子目录
      3. "02-{公司}-{系统}-保密承诺书.docx" 模板文件（从模板复制并替换公司名和日期）

    Args:
        project: 项目实体对象
        parent: 父级窗口（用于消息弹窗的模态绑定）
    """
    try:
        root = find_project_folder(project)
        if not root or not os.path.isdir(root):
            messagebox.showinfo("提示", "未找到项目文件夹")
            return
        cname = (project.company_name or "未命名").replace("/", "_").replace("\\", "_")
        sname = (project.system_name or "").replace("/", "_").replace("\\", "_")
        created = []; existed = []
        # 01-其他归档文件
        dname = "01-其他归档文件"
        dpath = os.path.join(root, dname)
        if not os.path.exists(dpath):
            os.makedirs(dpath, exist_ok=True); created.append(dname)
        else:
            existed.append(dname)
        # 00-报告打印
        dname = f"00-{cname}-{sname}-报告打印"
        dpath = os.path.join(root, dname)
        if not os.path.exists(dpath):
            os.makedirs(dpath, exist_ok=True); created.append(dname)
        else:
            existed.append(dname)
        # 02-保密承诺书
        nda_name = f"02-{cname}-{sname}-保密承诺书.docx"
        nda_path = os.path.join(root, nda_name)
        if os.path.exists(nda_path):
            existed.append(nda_name)
        else:
            try:
                tpl = os.path.join(Config.get_data_dir(), "templates", "02-保密承诺书模板.docx")
                if os.path.exists(tpl):
                    import docx
                    shutil.copy2(tpl, nda_path)
                    doc = docx.Document(nda_path)
                    company = project.company_name or ""
                    create_date = project.created_at[:10] if project.created_at else ""
                    # 替换公司名: 逐run保留格式
                    for p in doc.paragraphs:
                        for run in p.runs:
                            if "XX公司" in run.text or "xx公司" in run.text:
                                run.text = run.text.replace("XX公司", company).replace("xx公司", company)
                                break
                    # 清除 split 的 "XX"+"公司" run 对
                    for p in doc.paragraphs:
                        for j in range(len(p.runs) - 1):
                            if p.runs[j].text.strip() == "XX" and p.runs[j+1].text.strip() == "公司":
                                p.runs[j].text = company; p.runs[j+1].text = ""
                    # 替换日期: 保留每个run格式, 仅替换"XX"
                    if create_date:
                        parts = create_date.split("-")
                        if len(parts) == 3:
                            y, m, d = parts[0], f"{int(parts[1]):02d}", f"{int(parts[2]):02d}"
                            for table in doc.tables:
                                for row in table.rows:
                                    for cell in row.cells:
                                        for p in cell.paragraphs:
                                            rs = p.runs
                                            if len(rs) >= 6 and all(rs[k].text.strip() == v for k, v in [(0,"XX"),(2,"XX"),(4,"XX")]):
                                                rs[0].text = y; rs[2].text = m; rs[4].text = d
                                                break
                    doc.save(nda_path)
                    created.append(nda_name)
            except Exception:
                pass
        # 弹窗报告
        lines = []
        if created:
            lines.append("--- 已创建 ---")
            lines.extend(f"  + {x}" for x in created)
        if existed:
            lines.append("--- 已存在 ---")
            lines.extend(f"  = {x}" for x in existed)
        messagebox.showinfo("初始化完成", "\n".join(lines) if lines else "无需初始化")
    except Exception as e:
        messagebox.showerror("错误", f"初始化失败: {e}")


# =============================================================================
# on_rename_click - 批量重命名功能
# =============================================================================

def on_rename_click(project: Project, parent=None):
    """处理"一键重命名"按钮点击 - 批量修正项目过程文件的命名

    此方法是系统的核心自动化功能之一，执行以下流程：

    1. 查找项目文件夹（按 folder_path 或关键词搜索）
    2. 解压 ZIP 文件并重命名内容：
       - "测评方案评审记录表.zip" -> 解压提取文件
       - "测评报告评审表.zip" -> 解压提取文件（保留终审，删除初审）
    3. 删除包含"初审"的文件
    4. 修正文件命名格式为：{编号}-{公司}-{系统}-{标准名称}.{扩展名}
       - 已带编号的文件（如 "07-测评方案.docx"）：匹配关键词后更正编号和名称
       - 无编号的文件：自动添加编号和名称前缀
    5. 修正子目录命名格式：
       - "报告打印" -> "00-{公司}-{系统}-报告打印"
       - "渗透测试报告" -> "13-{公司}-{系统}-渗透测试报告"
    6. 输出操作报告（显示处理了多少文件、跳过多少文件）

    关键字映射表（key_map）：
      {匹配关键词: (编号, 标准化名称)}，按关键词长度倒序匹配避免误匹配

    Args:
        project: 项目实体对象
        parent: 父级窗口（用于消息弹窗的模态绑定）
    """
    try:
        root = find_project_folder(project)  # 查找项目文件夹路径
        if not root or not os.path.isdir(root):  # 文件夹不存在
            messagebox.showinfo("提示", "未找到项目文件夹")
            return

        # 生成文件名的前导前缀
        cname = (project.company_name or "未命名").replace("/", "_").replace("\\", "_")
        sname = (project.system_name or "").replace("/", "_").replace("\\", "_")
        new_prefix = f"{cname}-{sname}"  # 格式：公司名-系统名

        # ---- 关键字映射表 ----
        # 格式：{文件名关键词: (编号, 标准化显示名称)}
        # 按关键词长度倒序匹配，防止"测评方案评审表"误匹配为"测评方案"
        key_map = {
            "保密承诺书": ("02", "保密承诺书"),  # 02 号文件
            "测评调研表": ("03", "测评调研表"),  # 03 号文件
            "测评授权书": ("04", "测评授权书"),  # 04 号文件
            "风险告知书": ("05", "风险告知书"),  # 05 号文件
            "项目计划书": ("06", "项目计划书"),  # 06 号文件
            "测评方案": ("07", "测评方案"),  # 07 号文件（注意：需放在"测评方案评审记录表"之后匹配）
            "归档材料评审记录表": ("08", "测评方案评审表"),  # 08 号文件（历史名称）
            "测评方案评审表": ("08", "测评方案评审表"),  # 08 号文件（标准名称）
            "首次会议记录": ("09", "首次会议记录"),  # 09 号文件
            "测评现场记录表": ("10", "测评现场记录表"),  # 10 号文件
            "问题汇总": ("11", "问题汇总及整改建设书"),  # 11 号文件
            "漏洞扫描报告": ("12", "漏洞扫描报告"),  # 12 号文件
            "项目文档移交清单": ("14", "项目文档移交清单"),  # 14 号文件
            "末次会议记录": ("15", "末次会议记录"),  # 15 号文件
            "测评报告-终稿": ("16", "测评报告-终稿"),  # 16 号文件
            "测评报告评审记录表": ("17", "测评报告评审表"),  # 17 号文件（历史名称）
            "测评报告评审表": ("17", "测评报告评审表"),  # 17 号文件（标准名称）
            "服务情况评价表": ("18", "服务情况评价表"),  # 18 号文件
            "报备表": ("19", "报备表"),
            "渗透测试报告": ("13", "渗透测试报告"),  # 渗透测试报告目录
        }

        renamed = 0  # 重命名计数器（累计处理了多少个项目）
        msgs = []  # 操作报告消息列表（每步操作一条记录）

        # ========== 步骤 1: ZIP 文件解压处理 ==========
        for fname in os.listdir(root):  # 遍历项目根目录所有文件
            fpath = os.path.join(root, fname)  # 拼接文件完整路径
            if not os.path.isfile(fpath) or not fname.lower().endswith(".zip"):  # 跳过非 ZIP 文件
                continue

            # 处理"测评方案评审记录表.zip" -> 解压并重命名内容为 08
            if "测评方案评审记录表" in fname:
                try:
                    with zipfile.ZipFile(fpath, "r") as zf:  # 以读取模式打开 ZIP
                        zf.extractall(root)  # 解压全部内容到项目根目录
                    os.remove(fpath)  # 删除原 ZIP 文件
                    renamed += 1  # 计数 +1
                    msgs.append(f"解压: {fname} -> 提取文件")  # 记录操作
                except Exception as e:  # 解压失败（损坏的 ZIP、权限不足等）
                    msgs.append(f"解压失败: {fname} ({e})")

            # 处理"渗透测试报告.zip" -> 解压到独立文件夹后重命名
            if "渗透测试报告" in fname and "评审" not in fname:
                try:
                    # 解压到临时目录
                    tmp_dir = os.path.join(root, "_渗透测试报告_tmp")
                    if os.path.exists(tmp_dir):
                        shutil.rmtree(tmp_dir)
                    os.makedirs(tmp_dir)
                    with zipfile.ZipFile(fpath, "r") as zf:
                        zf.extractall(tmp_dir)
                    os.remove(fpath)
                    # 去除 ZIP 内的公共顶层目录前缀
                    items = os.listdir(tmp_dir)
                    if len(items) == 1 and os.path.isdir(os.path.join(tmp_dir, items[0])):
                        # ZIP自带一个顶层目录 -> 直接移动该目录
                        src = os.path.join(tmp_dir, items[0])
                        dst = os.path.join(root, f"13-{new_prefix}-渗透测试报告")
                    else:
                        # 文件散落 -> 创建目标目录，移入所有文件
                        dst = os.path.join(root, f"13-{new_prefix}-渗透测试报告")
                        if not os.path.exists(dst):
                            os.makedirs(dst)
                        for item in items:
                            shutil.move(os.path.join(tmp_dir, item), os.path.join(dst, item))
                        shutil.rmtree(tmp_dir)
                        renamed += 1
                        msgs.append(f"解压: {fname} -> 13-{new_prefix}-渗透测试报告/")
                        continue
                    if os.path.exists(dst):
                        shutil.rmtree(dst)
                    shutil.move(src, dst)
                    shutil.rmtree(tmp_dir)
                    renamed += 1
                    msgs.append(f"解压: {fname} -> 13-{new_prefix}-渗透测试报告/")
                except Exception as e:
                    msgs.append(f"解压失败: {fname} ({e})")

            # 处理"测评报告评审表.zip" -> 解压（后续会删除初审版本）
            elif "测评报告评审表" in fname:
                try:
                    with zipfile.ZipFile(fpath, "r") as zf:  # 以读取模式打开
                        zf.extractall(root)  # 解压到项目根目录
                    os.remove(fpath)  # 删除原 ZIP
                    renamed += 1
                    msgs.append(f"解压: {fname} -> 提取文件")
                except Exception as e:
                    msgs.append(f"解压失败: {fname} ({e})")

        # ========== 步骤 2: 删除初审版本文件 ==========
        for fname in os.listdir(root):
            if "初审" in fname:  # 包含"初审"的文件（如测评报告评审表-初审.docx）
                try:
                    os.remove(os.path.join(root, fname))  # 直接删除
                    msgs.append(f"删除初审: {fname}")
                except Exception:  # 删除失败（权限等），静默跳过
                    pass

        # ========== 步骤 3: 批量重命名文件 ==========
        for fname in os.listdir(root):  # 遍历根目录所有文件
            fpath = os.path.join(root, fname)
            if not os.path.isfile(fpath) or fname.endswith(".zip"):  # 跳过 ZIP 和目录
                continue
            name_no_ext, ext = os.path.splitext(fname)  # 分离文件名和扩展名（"07-测评方案", ".docx"）

            # ---- 情况 A: 文件名已有编号前缀（如 "07-测评方案.docx"） ----
            m = re.match(r"^(\d{2})-(.+)", name_no_ext)  # 匹配 "###-内容" 格式
            if m:  # 已有编号前缀
                num = m.group(1)  # 当前编号（如 "07"）
                rest = name_no_ext[len(num) + 1:]  # 编号后面的剩余部分（如 "测评方案.docx"）
                # 按关键词长度从大到小匹配，避免"测评方案评审表"被"测评方案"先匹配
                for keyword in sorted(key_map, key=len, reverse=True):
                    if keyword in rest:  # 剩余部分中包含关键字
                        num, target_kw = key_map[keyword]  # 获取正确的编号和标准名称
                        new_name = f"{num}-{new_prefix}-{target_kw}{ext}"  # 构建标准文件名
                        if new_name != fname:  # 需要重命名
                            new_path = os.path.join(root, new_name)
                            if not os.path.exists(new_path):  # 目标文件不存在
                                os.rename(fpath, new_path)  # 执行重命名
                                renamed += 1
                            else:  # 目标文件已存在（冲突）
                                msgs.append(f"跳过(已存在): {new_name}")
                        break  # 匹配到一个关键词后退出内层循环

            # ---- 情况 B: 文件名无编号（如 "测评调研表.docx"） ----
            else:
                for keyword in sorted(key_map, key=len, reverse=True):  # 按关键词长度倒序
                    if keyword in name_no_ext:  # 文件名中包含关键字
                        num, target_kw = key_map[keyword]  # 获取编号和标准名称
                        new_name = f"{num}-{new_prefix}-{target_kw}{ext}"  # 构建标准名
                        new_path = os.path.join(root, new_name)
                        if not os.path.exists(new_path):  # 不冲突
                            os.rename(fpath, new_path)  # 重命名
                            renamed += 1
                        else:
                            msgs.append(f"跳过(已存在): {new_name}")
                        break  # 匹配到一个后退出

        # ========== 步骤 4: 子目录重命名 ==========
        for dname in os.listdir(root):  # 遍历所有目录
            dpath = os.path.join(root, dname)
            if not os.path.isdir(dpath):  # 跳过文件
                continue
            for keyword, num in {"报告打印": "00", "渗透测试报告": "13"}.items():
                if keyword in dname and (cname not in dname or sname not in dname):  # 需要更新前缀
                    new_dname = f"{num}-{new_prefix}-{keyword}"  # 标准目录名
                    new_dpath = os.path.join(root, new_dname)
                    if not os.path.exists(new_dpath):  # 不冲突
                        os.rename(dpath, new_dpath)  # 重命名目录
                        renamed += 1
                    break  # 匹配到一个关键词后退出

        # ========== 步骤 5: 显示操作报告 ==========
        if msgs:  # 有详细操作记录
            msg_text = "\n".join(msgs[:15])  # 取前 15 条
            if len(msgs) > 15:  # 超过 15 条则显示统计
                msg_text += f"\n...共 {len(msgs)} 条"
            messagebox.showinfo("操作报告", msg_text)  # 弹窗展示
        elif renamed:  # 没有详细记录但统计 > 0
            messagebox.showinfo("完成", f"已处理 {renamed} 个项目")
        else:  # 无需任何修改
            messagebox.showinfo("提示", "所有文件名已是最新，无需修改")

    except Exception as e:  # 捕获所有未预见的异常
        messagebox.showerror("错误", f"操作失败: {e}")  # 弹窗显示错误


# =============================================================================
# on_zip_click - 打包过程文档功能
# =============================================================================

def on_zip_click(project: Project, parent=None):
    """处理"打包过程文档"按钮点击 - 将项目过程文件压缩为 ZIP

    打包策略：
      1. 查找项目文件夹
      2. 创建 ZIP 文件（命名格式：{公司}-{系统}-过程文档.zip）
      3. 按预定义的关键词列表筛选需要打包的文件：
         - 保密承诺书、测评调研表、测评授权书、风险告知书
         - 项目计划书、测评方案、首次会议记录、测评现场记录表
         - 问题汇总、漏洞扫描报告、项目文档移交清单、末次会议记录
         - 服务情况评价表、报备表
      4. 特殊处理"渗透测试报告"目录：
         - 如果目录非空：递归打包所有文件（保持目录结构）
         - 如果目录为空：添加空目录条目
      5. 成功打包后弹窗提示；无可打包文件时删除空 ZIP

    排除项：
      - 测评报告-终稿（不入过程文档包）
      - 报告打印相关文件
      - 其他归档文件

    Args:
        project: 项目实体对象
        parent: 父级窗口（用于消息弹窗的模态绑定）
    """
    try:
        root = find_project_folder(project)  # 查找项目文件夹路径
        if not root or not os.path.isdir(root):  # 文件夹不存在
            messagebox.showinfo("提示", "未找到项目文件夹")
            return

        # 构建 ZIP 文件名
        cname = project.company_name or "未命名"  # 公司名（取不到用"未命名"）
        sname = project.system_name or ""  # 系统名
        zip_name = f"{cname}-{sname}-过程文档.zip"  # ZIP 文件名格式
        zip_path = os.path.join(root, zip_name)  # ZIP 文件完整路径

        # ---- 需要打包的文件关键词列表 ----
        # 文件名中包含这些关键词之一的文件将被包含在 ZIP 中
        pack_keywords = [  # 仅打包 #3-#7 和 #9-#15
            "测评调研表", "测评授权书", "风险告知书",
            "项目计划书", "测评方案",
            "首次会议记录", "测评现场记录表",
            "问题汇总", "漏洞扫描报告",
            "项目文档移交清单", "末次会议记录",
        ]

        count = 0  # 已打包文件/目录计数
        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:  # 创建 ZIP（DEFLATED 压缩）
            # --- 第一步：打包匹配关键词的单文件 ---
            for fname in os.listdir(root):  # 遍历根目录
                fpath = os.path.join(root, fname)
                if not os.path.isfile(fpath) or fname == zip_name:  # 跳过目录和正在创建的 ZIP
                    continue
                name_no_ext = os.path.splitext(fname)[0]  # 取文件名（不含扩展名）
                for kw in pack_keywords:  # 遍历关键词
                    if kw in name_no_ext:  # 文件名匹配关键词
                        zf.write(fpath, fname)  # 写入 ZIP（保持原文件名）
                        count += 1  # 计数 +1
                        break  # 匹配到一个关键词即处理下一个文件

            # --- 第二步：打包渗透测试报告目录（含所有内容） ---
            for dname in os.listdir(root):
                dpath = os.path.join(root, dname)
                if os.path.isdir(dpath) and "渗透测试报告" in dname:  # 是目标目录
                    has_files = False  # 目录是否有文件（非空标志）
                    for dirpath, _, filenames in os.walk(dpath):  # 递归遍历目录
                        for fn in filenames:  # 遍历每个文件
                            fp = os.path.join(dirpath, fn)  # 文件完整路径
                            arcname = os.path.relpath(fp, root).replace("\\", "/")  # ZIP 内路径（统一斜杠）
                            zf.write(fp, arcname)  # 写入 ZIP，保持目录结构
                            count += 1
                            has_files = True  # 标记为非空
                    # 如果目录为空，添加空目录条目（保留目录结构）
                    if not has_files:
                        info = zipfile.ZipInfo(dname + "/")  # 创建目录条目（末尾 / 标记为目录）
                        zf.writestr(info, "")  # 写入空内容
                        count += 1

        # ---- 结果处理 ----
        if count > 0:  # 至少打包了一个文件
            messagebox.showinfo("打包完成",
                f"已打包 {count} 个文件\n{zip_name}")  # 显示打包结果
        else:  # 没有匹配的文件
            os.remove(zip_path)  # 删除空 ZIP 文件
            messagebox.showinfo("提示", "未找到可打包的过程文件")

    except Exception as e:  # 捕获所有异常
        messagebox.showerror("错误", f"打包失败: {e}")  # 弹窗显示错误


# =============================================================================
# on_report_print_click - 报告打印功能
# =============================================================================

def on_report_print_click(project: Project, parent=None):
    """处理"报告打印"按钮点击 - 生成测评报告打印信息并复制相关文件

    完整流程：
      1. 查找项目文件夹（如果找不到则提示并退出）
      2. 弹出 show_report_dialog 编辑确认框，预填项目数据
      3. 用户确认 14 个字段后，查找/创建报告打印子目录
      4. 调用 _create_report_xlsx_data 生成测评报告打印信息.xlsx
      5. 复制相关文件到报告打印目录：
         - "测评授权书"相关文件
         - "风险告知书"相关文件
         - "测评报告-终稿.pdf"（如果有）
         - 过程文档 ZIP（如果有）
      6. 弹窗显示生成结果

    报告打印目录命名：00-{公司}-{系统}-报告打印
    XLSX 文件命名：00-{公司}-{系统}-测评报告打印信息.xlsx

    Args:
        project: 项目实体对象
        parent: 父级窗口（用于对话框的模态绑定）
    """
    try:
        proot = find_project_folder(project)  # 查找项目文件夹根目录
        if not proot or not os.path.isdir(proot):  # 根目录不存在
            messagebox.showinfo("提示", "未找到项目文件夹")
            return

        # ---- 步骤 1: 弹出报告打印编辑确认框 ----
        data = show_report_dialog(
            parent,  # 父窗口
            cname=project.company_name or "",  # 预填公司名称
            sname=project.system_name or "",  # 预填系统名称
            location=(project.location or "").split("-")[0]  # 预填所属地（取省/市前缀）
                if project.location else "",  # 无属地则空
            deadline=project.deadline or date.today().strftime("%Y-%m-%d"),  # 预填日期
        )
        if not data:  # 用户点击了取消
            return

        # 提取用户确认后的关键字段
        cname = data["cname"]  # 客户公司全称
        sname = data["sname"]  # 系统名称
        prefix = f"{cname}-{sname}"  # 文件名前缀

        # ---- 步骤 2: 查找或创建报告打印目录 ----
        report_dir = None  # 报告打印目录路径
        for dname in os.listdir(proot):  # 在项目根目录中查找
            if "报告打印" in dname and os.path.isdir(os.path.join(proot, dname)):
                report_dir = os.path.join(proot, dname)  # 找到已存在的报告打印目录
                break
        if not report_dir:  # 不存在则创建
            report_dir = os.path.join(proot, f"00-{prefix}-报告打印")  # 按标准格式命名
            os.makedirs(report_dir, exist_ok=True)  # 递归创建

        # ---- 步骤 3: 创建测评报告打印信息 XLSX ----
        xlsx_name = f"00-{prefix}-测评报告打印信息.xlsx"  # XLSX 文件名
        xlsx_path = os.path.join(report_dir, xlsx_name)  # XLSX 完整路径
        _create_report_xlsx_data(project, xlsx_path, data, report_dir, proot)  # 生成 XLSX

        # ---- 步骤 4: 复制相关文件到报告打印目录 ----
        copied = 0  # 已复制文件计数

        # 复制"测评授权书"和"风险告知书"相关文件
        copy_keywords = ["测评授权书", "风险告知书"]
        for fname in os.listdir(proot):  # 遍历项目根目录
            fpath = os.path.join(proot, fname)
            if not os.path.isfile(fpath):  # 跳过目录
                continue
            for kw in copy_keywords:  # 遍历复制关键词
                if kw in fname:  # 文件名匹配关键词
                    shutil.copy2(fpath, os.path.join(report_dir, fname))  # 复制（保留元数据）
                    copied += 1  # 计数
                    break  # 匹配一个即处理下一个文件

            # 复制测评报告终稿 PDF
            if "测评报告-终稿" in fname and fname.lower().endswith(".pdf"):  # 必须是 PDF
                shutil.copy2(fpath, os.path.join(report_dir, fname))  # 复制到打印目录
                copied += 1

        # 复制过程文档 ZIP（如果已生成）
        zip_name = f"{cname}-{sname}-过程文档.zip"  # ZIP 文件名
        zip_src = os.path.join(proot, zip_name)  # ZIP 源路径
        if os.path.exists(zip_src):  # ZIP 文件存在
            shutil.copy2(zip_src, os.path.join(report_dir, zip_name))  # 复制
            copied += 1

        # ---- 步骤 5: 显示结果 ----
        messagebox.showinfo("报告打印完成",
            f"已生成 {xlsx_name}\n已复制 {copied} 个文件到报告打印目录")

    except Exception as e:  # 捕获所有异常
        messagebox.showerror("错误", f"报告打印失败: {e}")  # 显示错误
