"""
项目初始化模块 -- 等保测评进度管理系统

============ 本模块职责 ============
为新建项目创建标准化的本地文件结构：
  {序号}-{公司}-{系统}-{YYMMDD}/
    ├── 01-其他归档文件/
    │   ├── 00-网安报备
    │   ├── 01-备案材料
    │   ├── 02-往期测评报告
    │   ├── 03-现场测评
    │   └── 04-渗透漏扫
    └── 02-{公司}-{系统}-保密承诺书.docx  (基于 docx 模板自动生成)

============ 命名规范 ============
  - 保密承诺书（多系统合并项目）：02-{公司名称}-保密承诺书.docx
  - 保密承诺书（单系统项目）：02-{公司名称}-{系统名称}-保密承诺书.docx

============ docx 模板替换逻辑 ============
保密承诺书基于 templates/02-保密承诺书模板.docx 模板生成，
替换其中的占位符：
  1. 公司名占位符："XX公司" 或 "xx公司" → 实际公司名
  2. 分体占位符："XX" + "公司" 两个相邻 run → 合并替换为公司名
  3. 日期占位符：表格中第1/3/5个 run 为 "XX" → 年月日
"""

import os        # 操作系统接口：路径拼接、目录创建、文件存在检查
import shutil    # 文件操作工具：复制模板文件到目标位置
from tkinter import messagebox  # 提示框：显示初始化结果

from models.project import Project               # 项目实体类
from utils.config import Config                   # 全局配置：获取模板文件路径
from ui.file_ops.folder_ops import find_project_folder  # 查找项目文件夹


# =============================================================================
# on_init_click - 项目初始化入口函数
# =============================================================================

def on_init_click(project: Project, parent=None, all_projects: list = None):
    """执行项目初始化：创建标准子目录和保密承诺书模板文件。

    完整的初始化流程（4 个阶段）：
      阶段 1：定位项目文件夹 → 调用 find_project_folder 查找
      阶段 2：创建归档子目录 → "01-其他归档文件" 文件夹
      阶段 3：生成保密承诺书 → 基于模板替换占位符
      阶段 4：弹窗报告结果 → 汇总已创建 / 已存在的项目

    保密承诺书模板替换逻辑（docx 文档的 XML 结构）：
      - 每个段落（paragraph）由多个 run 组成，run 是保留格式的最小单元
      - Word 中如果某个词中间修改了格式，该词会被拆成多个 run
      - 本函数需要处理两种占位符形态：
        形态 A：整个 "XX公司" 在同一个 run 中 → 直接字符串替换
        形态 B："XX" + "公司" 被拆成两个相邻 run（格式不同导致）→ 分别替换

    Args:
        project: 项目实体对象，包含公司名称、系统名称、创建日期等信息。
        parent: 父级 Tkinter 窗口引用（用于模态对话框的 parent 参数）。
        all_projects: 多系统合并时所有关联项目的列表，用于判断命名模式。
                      为 None 或长度 ≤ 1 时按单系统模式处理。

    Returns:
        None: 所有结果通过 messagebox 弹窗通知用户。
    """
    try:
        root = find_project_folder(project)
        if not root or not os.path.isdir(root):
            messagebox.showinfo("提示", "未找到项目文件夹")
            return

        cname = (project.company_name or "未命名").replace("/", "_").replace("\\", "_")
        sname = (project.system_name or "").replace("/", "_").replace("\\", "_")
        is_multi = all_projects and len(all_projects) > 1
        created = []
        existed = []

        # ========== 阶段 2：创建归档子目录 ==========
        archive_root = os.path.join(root, "01-其他归档文件")
        archive_subdirs = [
            "00-网安报备", "01-备案材料", "02-往期测评报告",
            "03-现场测评", "04-渗透漏扫",
        ]
        for dname in archive_subdirs:
            dpath = os.path.join(archive_root, dname)
            if not os.path.exists(dpath):
                os.makedirs(dpath, exist_ok=True)  # 创建不存在的目录
                created.append(f"01-其他归档文件/{dname}")
            else:
                existed.append(f"01-其他归档文件/{dname}")

        # ========== 阶段 2.5：多系统(2+)创建子系统目录 ==========
        if is_multi:
            for p in all_projects:
                sn = (p.system_name or "").replace("/", "_").replace("\\", "_")
                if sn:
                    dpath = os.path.join(root, sn)
                    if not os.path.exists(dpath):
                        os.makedirs(dpath, exist_ok=True)
                        created.append(sn)
                    else:
                        existed.append(sn)

        # ========== 阶段 3：生成保密承诺书 docx ==========
        nda_name = f"02-{cname}-保密承诺书.docx" if is_multi else f"02-{cname}-{sname}-保密承诺书.docx"
        nda_path = os.path.join(root, nda_name)

        if os.path.exists(nda_path):
            # 保密承诺书已存在则不重复生成，避免覆盖用户已修改的版本
            existed.append(nda_name)
        else:
            try:
                tpl = os.path.join(Config.get_data_dir(), "templates", "02-保密承诺书模板.docx")
                if os.path.exists(tpl):
                    import docx  # python-docx 库：读写 Word 文档（延迟导入避免循环依赖）
                    # 复制模板到目标位置（shutil.copy2 保留文件元数据，如修改时间）
                    shutil.copy2(tpl, nda_path)
                    doc = docx.Document(nda_path)  # 打开刚复制的文档准备编辑

                    company = project.company_name or ""  # 当前项目的公司名称（用于替换占位符）
                    # 提取创建日期并截取前 10 位（YYYY-MM-DD 格式），用于日期列替换
                    create_date = project.created_at[:10] if project.created_at else ""

                    # ---- 替换公司名（形态 A：完整 "XX公司" 在同一个 run 中） ----
                    # 遍历所有段落的每个 run，直接进行字符串替换。
                    # Python-docx 中 run 是保留格式的最小单元，这样替换后
                    # 原有的字体、字号、颜色等格式属性都得以保留。
                    for p in doc.paragraphs:
                        for run in p.runs:
                            if "XX公司" in run.text or "xx公司" in run.text:
                                # 同时替换大写和小写版本，兼容模板编辑时可能的不同输入法习惯
                                run.text = run.text.replace("XX公司", company).replace("xx公司", company)
                                break  # 一个段落中通常只有一个占位符，找到即跳出

                    # ---- 替换公司名（形态 B：分体的 "XX" + "公司" 相邻 run） ----
                    # 某些模板中 "XX" 和 "公司" 因格式不同（如宋体 vs 黑体）被 Word
                    # 拆分成两个相邻的 run，例如：run[0]="XX"(宋体) + run[1]="公司"(黑体)。
                    # 此时需要将 run[0] 改为完整的公司名，run[1] 清空以避免文字重复。
                    for p in doc.paragraphs:
                        for j in range(len(p.runs) - 1):  # 遍历相邻 run 对
                            if p.runs[j].text.strip() == "XX" and p.runs[j+1].text.strip() == "公司":
                                p.runs[j].text = company   # 第一个 run 填入完整公司名
                                p.runs[j+1].text = ""      # 第二个 run 清空（避免出现"公司"残留）

                    # ---- 替换日期占位符（表格中的年-月-日） ----
                    # 模板表格中的日期行格式：[XX]年[XX]月[XX]日，
                    # 每个 "XX" 在独立的 run 中（保留各自的对齐和字体格式）。
                    if create_date:
                        parts = create_date.split("-")  # "2026-06-12" → ["2026", "06", "12"]
                        if len(parts) == 3:
                            # 月份和日期去除前导零后重新补零格式化
                            y, m, d = parts[0], f"{int(parts[1]):02d}", f"{int(parts[2]):02d}"
                            # 遍历文档中所有表格（含页眉页脚表格）的每个单元格
                            for table in doc.tables:
                                for row in table.rows:
                                    for cell in row.cells:
                                        for p in cell.paragraphs:
                                            rs = p.runs  # 当前段落的所有 run 列表
                                            # 匹配 6 个 run 中第 0/2/4 位均为 "XX" 的典型日期行
                                            # 模板结构：[XX]  年  [XX]  月  [XX]  日
                                            if len(rs) >= 6 and all(
                                                rs[k].text.strip() == v
                                                for k, v in [(0, "XX"), (2, "XX"), (4, "XX")]
                                            ):
                                                rs[0].text = y   # run[0] 替换为年份
                                                rs[2].text = m   # run[2] 替换为月份
                                                rs[4].text = d   # run[4] 替换为日期
                                                break  # 一个单元格中通常只有一个日期行

                    # 保存修改后的文档（commit 到 docx 文件）
                    doc.save(nda_path)
                    created.append(nda_name)
            except Exception:
                import traceback
                print(f"[初始化] NDA创建失败:", flush=True)
                traceback.print_exc()

        # ========== 阶段 4：弹窗汇总报告 ==========
        lines = []
        if created:
            lines.append("--- 已创建 ---")
            lines.extend(f"  + {x}" for x in created)
        if existed:
            lines.append("--- 已存在 ---")
            lines.extend(f"  = {x}" for x in existed)
        messagebox.showinfo("初始化完成", "\n".join(lines) if lines else "无需初始化")

    except Exception as e:
        import traceback
        traceback.print_exc()
        messagebox.showerror("错误", f"初始化失败: {e}")
