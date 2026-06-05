"""
主窗口模块 - 应用程序的主界面控制器

本模块是整个"等保测评进度管理系统"的核心控制器（Controller），遵循 MVC 设计模式：
  - Model（模型层）：models 包中的 Project、WorkflowStage 等实体类
  - View（视图层）：ui 包中的 Toolbar、KanbanBoard、ProjectCard 等组件
  - Controller（控制层）：MainWindow 自身，负责协调视图与模型的交互

主要职责：
  1. 工具栏、看板等 UI 组件的创建和布局管理
  2. 业务服务层（数据、项目、流程、日志）的初始化与依赖注入
  3. 用户交互事件（按钮点击、卡片操作、窗口事件）的响应与分发
  4. 窗口生命周期管理（创建时加载数据、关闭时自动保存）

技术栈：
  - Tkinter：Python 标准 GUI 库，提供窗口、Frame、组件等
  - 服务层注入：DataService / ProjectService / WorkflowService / LogService
"""

# =============================================================================
# 导入区 - 按层级组织
# =============================================================================

import tkinter as tk  # Python 标准 GUI 库，用于构建桌面端窗口和容器组件
from tkinter import messagebox  # 消息弹窗组件，用于显示提示、警告、错误和信息确认对话框

# ---- 模型层 ----
from models.workflow import WorkflowStage  # 流程阶段实体类，表示看板中的一列（如"项目启动"）

# ---- 服务层 ----
from services.data_service import DataService  # 数据持久化服务：负责 JSON 文件的读写操作
from services.project_service import ProjectService  # 项目业务服务：处理项目的增删改查和阶段移动
from services.workflow_service import WorkflowService  # 流程业务服务：管理流程阶段的配置和宽度
from services.log_service import LogService  # 日志服务：记录和查询所有操作日志

# ---- 视图层 ----
from ui.toolbar import Toolbar  # 顶部工具栏组件，包含新增、编辑、删除等快捷按钮
from ui.kanban_board import KanbanBoard  # 看板容器组件，承载所有流程阶段列和项目卡片
from ui.project_card import ProjectCard  # 项目卡片组件（仅用于类型标注和回调参数类型提示）
from ui.project_dialog import show_project_dialog  # 项目编辑对话框便捷函数（新增/编辑项目表单）
from ui.workflow_dialog import WorkflowDialog  # 流程编辑对话框类，管理流程阶段的增删改排序
from ui.log_dialog import show_log_dialog  # 日志查看对话框便捷函数（展示操作历史记录）
from ui.detail_dialog import show_detail_dialog  # 项目详情对话框便捷函数（展示项目完整信息）
from ui.backup_dialog import BackupDialog  # WebDAV 备份对话框类，负责远程备份管理

# ---- 工具层 ----
from utils.config import Config  # 全局配置类，提供应用名称、版本、窗口尺寸、颜色等常量


class MainWindow(tk.Tk):
    """应用程序主窗口 - 继承自 tk.Tk（Tkinter 顶层窗口）

    作为整个应用的控制器，MainWindow 负责串联所有子系统：
      - 创建并注入服务层依赖（DataService、ProjectService、WorkflowService、LogService）
      - 构建并布局 UI 组件（Toolbar 工具栏 + KanbanBoard 看板）
      - 绑定用户交互事件的回调函数（按钮点击、卡片操作、窗口调整）
      - 管理窗口生命周期（初始化加载数据、关闭前保存数据）

    设计原则：
      - 单一数据源：所有业务数据通过服务层管理，UI 仅负责展示
      - 事件驱动：UI 组件通过回调函数通知控制器处理业务逻辑
      - 依赖注入：服务层实例在 __init__ 中创建并传递给需要的组件

    Attributes:
        _data_service (DataService): 数据持久化服务实例
        _log_service (LogService): 操作日志服务实例
        _project_service (ProjectService): 项目业务服务实例
        _workflow_service (WorkflowService): 流程业务服务实例
        _toolbar (Toolbar): 顶部工具栏组件
        _kanban (KanbanBoard): 主看板组件
    """

    def __init__(self):
        """初始化主窗口及所有子系统

        初始化顺序严格按照依赖关系执行：
          1. 创建服务层实例（DataService -> LogService -> ProjectService -> WorkflowService）
          2. 配置窗口属性（标题、尺寸、最小尺寸、背景色）
          3. 构建 UI 组件（先工具栏、后看板）
          4. 绑定窗口事件（大小变化监听、关闭协议钩子）
          5. 加载数据并刷新看板显示
        """
        super().__init__()  # 调用父类 tk.Tk 的构造方法，创建顶级窗口

        # =====================================================================
        # 第一步：初始化服务层（数据 -> 日志 -> 项目 -> 流程，按依赖顺序创建）
        # =====================================================================

        # --- 数据服务：处理 JSON 文件的读写，是所有数据持久化的入口 ---
        data_file = Config.get_data_file_path()  # 获取主数据文件完整路径（如 ./data/dap_data.json）
        self._data_service = DataService(data_file)  # 创建数据服务实例并传入文件路径

        # --- 日志服务：记录和查询所有用户操作日志（增删改、阶段移动等） ---
        self._log_service = LogService()  # 创建日志服务实例（内部自行管理日志文件路径）

        # --- 项目服务：封装所有项目相关的业务逻辑 ---
        self._project_service = ProjectService(
            self._data_service,  # 注入数据服务：项目操作需要读写 JSON 文件
            log_callback=self._log_service.create_log_callback(),  # 注入日志回调：每个项目操作自动记录日志
        )

        # --- 流程服务：管理流程阶段（列）的配置，包括阶段颜色、宽度、顺序等 ---
        self._workflow_service = WorkflowService(
            self._data_service,  # 注入数据服务：流程阶段数据同样保存在 JSON 文件中
            log_callback=self._log_service.create_log_callback(),  # 注入日志回调：流程编辑操作自动记录
        )

        # =====================================================================
        # 第二步：配置窗口基本属性
        # =====================================================================

        self.title(f"{Config.APP_NAME} v{Config.APP_VERSION}")  # 设置窗口标题：应用名称 + 版本号
        self.geometry(f"{Config.WINDOW_WIDTH}x{Config.WINDOW_HEIGHT}")  # 设置初始窗口尺寸（宽度x高度）
        self.minsize(Config.WINDOW_MIN_WIDTH, Config.WINDOW_MIN_HEIGHT)  # 设置最小窗口尺寸，防止缩得过小
        self.configure(bg=Config.KANBAN_BG)  # 设置窗口默认背景色（与看板背景统一，浅灰蓝色）

        # =====================================================================
        # 第三步：构建 UI 组件（从上到下：工具栏 -> 看板）
        # =====================================================================

        self._build_toolbar()  # 构建顶部工具栏区域（包含新增、编辑、删除、刷新等按钮）
        self._build_kanban()  # 构建主看板区域（包含流程阶段列和项目卡片，填充剩余空间）

        # =====================================================================
        # 第四步：绑定窗口级别的全局事件
        # =====================================================================

        self.bind("<Configure>", self._on_window_resize)  # 窗口大小变化时触发（预留扩展点）
        self.protocol("WM_DELETE_WINDOW", self._on_close)  # 点击关闭按钮时，先保存数据再退出

        # =====================================================================
        # 第五步：加载数据并刷新看板（从服务层读取阶段和项目数据，渲染到 UI）
        # =====================================================================

        self._refresh_kanban()
        # 启动后延迟检测云端备份（等 UI 完全加载）
        self.after(1000, self._check_restore_on_startup)

    # ==================================================================================
    # UI 构建方法 - 创建和布局所有界面组件
    # ==================================================================================

    def _build_toolbar(self):
        """构建顶部工具栏组件

        创建 Toolbar 实例并固定在窗口顶部（pack 布局，top 侧对齐，水平填充）。
        设置工具栏各个按钮的回调函数，将按钮点击事件连接到对应的业务处理方法。

        工具栏按钮对应关系：
          - "新增项目"  -> _on_add_project()    # 打开新增项目对话框
          - "编辑流程"  -> _on_edit_workflow()   # 打开流程阶段配置对话框
          - "查看日志"  -> _on_view_logs()       # 打开操作日志查看窗口
          - "删除项目"  -> _on_delete_selected() # 删除当前选中的项目
          - "刷新"      -> _refresh_kanban()     # 重新加载数据并刷新看板
          - "WebDAV备份" -> _on_backup()         # 打开 WebDAV 备份管理对话框
        """
        self._toolbar = Toolbar(self)  # 创建工具栏组件实例，父容器为当前主窗口（self）
        self._toolbar.pack(fill=tk.X, side=tk.TOP)  # 固定在顶部，水平方向填充整个窗口宽度

        # 绑定工具栏各按钮的回调函数 - 将 UI 事件连接到 MainWindow 的业务处理方法
        self._toolbar.on_add_project = self._on_add_project  # "新增项目" 按钮 -> 新增项目处理
        self._toolbar.on_edit_workflow = self._on_edit_workflow  # "编辑流程" 按钮 -> 流程编辑处理
        self._toolbar.on_view_logs = self._on_view_logs  # "操作日志" 按钮 -> 查看日志处理
        self._toolbar.on_delete_project = self._on_delete_selected  # "删除项目" 按钮 -> 删除选中项目
        self._toolbar.on_refresh = self._refresh_kanban  # "刷新" 按钮 -> 刷新看板数据
        self._toolbar.on_backup = self._on_backup
        self._toolbar.on_console = self._on_console  # "控制台" 按钮

    def _build_kanban(self):
        """构建看板区域组件

        创建 KanbanBoard 实例并放置在工具栏下方（fill=BOTH + expand=True 使其填充所有剩余空间）。
        设置看板内各种用户交互事件的回调函数，将卡片操作连接到控制器的业务处理方法。

        看板交互事件对应关系：
          - 卡片单击      -> _on_card_selected()    # 选中/取消选中卡片
          - 卡片双击      -> _on_card_edit()        # 打开编辑对话框
          - 详情按钮      -> _on_card_detail()      # 打开项目详情窗口
          - 编辑按钮      -> _on_card_edit()        # 打开编辑对话框（同双击）
          - 箭头移动      -> _on_card_move_stage()  # 将项目移动到前/后阶段
          - 复制按钮      -> _on_card_copy()        # 复制当前项目
          - 列宽拖拽      -> _on_column_resize()    # 保存调整后的列宽
        """
        self._kanban = KanbanBoard(self)  # 创建看板组件实例，父容器为当前主窗口
        self._kanban.pack(fill=tk.BOTH, expand=True, side=tk.TOP)  # 填充剩余空间，可随窗口缩放

        # 绑定看板中各种卡片操作的回调函数
        self._kanban.on_card_click = self._on_card_selected  # 卡片单点 -> 选中/取消选中
        self._kanban.on_card_double_click = self._on_card_edit  # 卡片双击 -> 打开项目编辑对话框
        self._kanban.on_card_detail = self._on_card_detail  # "详情" 按钮 -> 打开项目详情窗口
        self._kanban.on_card_edit = self._on_card_edit  # "编辑" 按钮 -> 打开项目编辑对话框
        self._kanban.on_card_move_stage = self._on_card_move_stage  # 箭头按钮 -> 移动项目阶段
        self._kanban.on_card_copy = self._on_card_copy  # "复制" 按钮 -> 复制当前项目
        self._kanban.on_column_resize = self._on_column_resize  # 列宽拖拽 -> 保存新宽度

    # ==================================================================================
    # 工具栏事件处理方法 - 响应顶部工具栏各按钮的点击
    # ==================================================================================

    def _on_add_project(self):
        """处理"新增项目"按钮点击事件

        执行流程：
          1. 检查是否有已配置的流程阶段（无阶段时无法创建项目，提示用户先配置）
          2. 打开新增项目对话框，预填当前所有流程阶段供下拉选择
          3. 用户填写并确认后，通过 ProjectService 创建项目实体
          4. 创建成功后自动建立项目文件夹结构，刷新看板显示

        数据流：
          用户输入 -> show_project_dialog() 返回 dict -> ProjectService.create_project() -> 写入 JSON
        """
        stages = self._workflow_service.get_all_stages()  # 从流程服务获取所有已配置的阶段列表
        if not stages:  # 如果没有任何流程阶段
            messagebox.showwarning("提示", "请先配置流程阶段")  # 弹窗提示用户需要先配置阶段
            return  # 中止操作，不打开对话框

        result = show_project_dialog(self, "新增项目", stages=stages)  # 打开新增项目对话框，传入阶段列表
        if result:  # 用户点击了"确认"按钮（非取消），result 为包含所有表单字段的字典
            # 调用项目服务创建项目，传入用户在对话框中填写的各字段值
            success, msg, project = self._project_service.create_project(
                company_name=result["company_name"],  # 客户公司名称
                system_name=result["system_name"],  # 被测系统名称
                cert_number=result["cert_number"],  # 证书备案编号
                issue_date=result.get("issue_date", ""),  # 证书签发日期（可选）
                level=result.get("level", ""),  # 系统安全保护等级（可选）
                location=result.get("location", ""),  # 项目所属地（可选）
                deadline=result["deadline"],  # 项目截止交付日期（必填）
                notes=result["notes"],  # 备注信息
                stage_id=result["stage_id"],  # 初始流程阶段 ID
            )
            if success:
                import os
                user_folder = result.get("folder_path", "").strip()
                if user_folder and os.path.isdir(user_folder):
                    # 用户指定了已有目录，直接使用
                    project.folder_path = user_folder
                    self._data_service.update_project(project.id, {"folder_path": user_folder})
                else:
                    self._create_project_folder(project)
                self._refresh_kanban()
            else:
                messagebox.showerror("错误", msg)  # 创建失败时弹窗显示错误原因

    def _on_edit_workflow(self):
        """处理"编辑流程"按钮点击事件

        打开 WorkflowDialog 流程编辑器，允许用户进行以下操作：
          - 新增流程阶段（填写名称、选择颜色）
          - 删除现有阶段（移除不需要的步骤）
          - 重命名阶段名称
          - 调整阶段之间的排列顺序（上移/下移）

        保存后的容错处理：
          - 如果某个阶段被删除，则该阶段下的所有项目自动迁移到新配置的第一个阶段
          - 避免出现"孤儿项目"（项目归属于不存在的阶段）

        操作完成后记录日志并刷新看板。
        """
        old_stages = self._workflow_service.get_all_stages()  # 获取修改前的阶段列表（用于对比哪些阶段被删除）
        dialog = WorkflowDialog(self, old_stages)  # 打开流程编辑对话框，传入当前阶段列表
        self.wait_window(dialog)  # 等待对话框关闭（模态阻塞，用户操作期间主窗口不可交互）

        if dialog.result:  # 用户点击了"保存"按钮，dialog.result 为编辑后的阶段配置列表
            # 将新的流程配置保存到数据服务（直接替换所有阶段数据）
            self._data_service.replace_all_stages(dialog.result)

            # 检测被删除的阶段：计算新旧阶段 ID 集合的差集
            new_stage_ids = {s["id"] for s in dialog.result}  # 新配置中的所有阶段 ID（转为集合）
            old_stage_ids = {s.id for s in old_stages}  # 旧配置中的所有阶段 ID（转为集合）
            removed_ids = old_stage_ids - new_stage_ids  # 差集 = 被删除的阶段 ID

            if removed_ids and new_stage_ids:  # 确实有阶段被删除，且新配置中至少还保留一个阶段
                first_new_id = dialog.result[0]["id"]  # 获取新配置的第一个阶段 ID（作为兜底目标）
                for pid in removed_ids:  # 遍历每一个被删除的阶段 ID
                    # 查找所有归属于被删除阶段的项目
                    for proj in self._project_service.get_all_projects():
                        if proj.stage_id == pid:  # 发现一个"孤儿项目"
                            # 将该项目的阶段更新为新配置的第一个阶段
                            self._project_service.update_project(
                                proj.id,  # 项目唯一 ID
                                stage_id=first_new_id,  # 迁移到第一个阶段
                            )

            # 记录本次编辑流程的操作日志
            self._log_service.add_log(
                action="编辑流程",  # 操作类型
                detail="更新流程阶段配置",  # 操作详细描述
            )
            self._refresh_kanban()  # 刷新看板显示（反映阶段配置的最新变化）

    def _on_view_logs(self):
        """处理"查看日志"按钮点击事件

        从日志服务获取所有系统操作记录，然后打开日志查看对话框展示给用户。
        日志按时间倒序排列，最新记录显示在最前面。
        """
        logs = self._log_service.get_all_logs()  # 获取所有操作日志列表（已按时间倒序排列）
        show_log_dialog(self, logs)  # 打开日志查看对话框，展示日志内容

    def _on_console(self):
        """打开控制台窗口，查看程序错误日志"""
        import tkinter as tk
        from tkinter import scrolledtext
        from utils.error_log import get_errors
        dlg = tk.Toplevel(self)
        dlg.title("控制台 - 错误日志")
        dlg.geometry("700x400")
        dlg.configure(bg="#1e1e1e")
        text = scrolledtext.ScrolledText(dlg, bg="#1e1e1e", fg="#d4d4d4",
            insertbackground="white", font=("Consolas", 9), wrap="word")
        text.pack(fill=tk.BOTH, expand=True, padx=4, pady=4)
        text.insert("1.0", get_errors())
        text.configure(state="disabled")
        btn = tk.Button(dlg, text="关闭", command=dlg.destroy,
            bg="#333", fg="#ccc", cursor="hand2")
        btn.pack(pady=(0, 4))

    def _on_backup(self):
        """处理"WebDAV备份"按钮点击事件"""
        dialog = BackupDialog(self, Config.get_data_file_path())  # 创建备份对话框，传入数据文件路径
        # 设置恢复回调：当用户执行恢复操作后，重新加载数据文件并刷新看板
        dialog.on_restore = lambda: (self._data_service.reload(), self._refresh_kanban())
        self.wait_window(dialog)  # 等待备份对话框关闭

    def _on_delete_selected(self):
        """处理"删除项目"按钮点击事件

        执行流程：
          1. 获取当前看板中选中的项目 ID
          2. 如果未选中任何卡片，提示用户先选择项目
          3. 显示二次确认对话框（防止误删）
          4. 确认后通过 ProjectService 执行永久删除
          5. 删除成功后刷新看板
        """
        project_id = self._kanban.get_selected_project_id()  # 从看板获取当前选中卡片的项目 ID
        if not project_id:  # 没有任何卡片被选中
            messagebox.showinfo("提示", "请先在卡片上点击选择要删除的项目")  # 提示用户先选择
            return  # 中止操作

        project = self._project_service.get_project_by_id(project_id)  # 根据 ID 获取项目对象（用于显示名称）
        if not project:  # 防御性检查：项目可能已被其他操作删除
            return

        # 二次确认对话框：显示项目名称，提醒用户此操作不可撤销
        if messagebox.askyesno("确认删除",
                               f"确定要永久删除项目\u300c{project.name}\u300d吗？\n\n"
                               "此操作不可撤销！"):  # \u300c = 「，\u300d = 」
            success, msg = self._project_service.delete_project(project_id)  # 执行项目删除操作
            if success:  # 删除成功
                self._refresh_kanban()  # 刷新看板，移除已删除的项目卡片
            else:
                messagebox.showerror("错误", msg)  # 删除失败时弹窗显示错误信息

    # ==================================================================================
    # 看板事件处理方法 - 响应看板中卡片的交互操作
    # ==================================================================================

    def _on_card_selected(self, card: ProjectCard | None):
        """处理卡片选中状态变化事件

        当用户单击看板中的卡片时触发。可用于扩展功能，如：
          - 更新状态栏显示当前选中项目的信息
          - 启用/禁用依赖于选中状态的操作按钮

        当前为空实现，预留扩展点。

        Args:
            card: 当前选中的卡片组件实例，或 None（表示取消选中/点击空白区域）
        """
        pass  # 预留扩展：可在此添加状态栏更新、右键菜单等逻辑

    def _on_card_detail(self, card: ProjectCard):
        """处理卡片"详情"按钮点击事件 - 打开项目详情查看窗口

        展示项目的完整信息，包括：
          - 基本信息（公司名称、系统名称、等级、地点等）
          - 流程阶段信息（当前所处阶段）
          - 操作日志（该项目的增删改历史记录）
          - 下级文件（项目目录中的过程文档列表）

        在详情窗口中支持的操作：
          - 编辑项目信息（edit） -> _handle_edit 内部闭包
          - 删除项目（delete） -> _handle_delete 内部闭包
          - 阶段移动（move）   -> _handle_move 内部闭包

        内部闭包说明：
          _handle_move: 在详情窗口中移动项目阶段后，重新获取最新数据刷新详情显示，
                        避免关闭窗口后重新打开的麻烦。

        Args:
            card: 被点击"详情"按钮的项目卡片组件
        """
        stages = self._workflow_service.get_all_stages()  # 获取当前所有流程阶段（供阶段下拉选择）
        logs = self._log_service.get_project_logs(card.project.id)  # 获取该项目的专属操作日志

        def _handle_move(target_stage_id, dialog):
            """内部闭包：在详情窗口中移动项目阶段，并刷新详情窗口数据

            Args:
                target_stage_id: 目标阶段的唯一标识符
                dialog: 详情对话框实例（用于刷新数据）
            """
            success, msg = self._project_service.move_project(
                card.project.id, target_stage_id)  # 调用项目服务执行阶段移动
            if success:  # 移动成功
                self._refresh_kanban()  # 刷新看板，同步卡片位置
                upd = self._project_service.get_project_by_id(card.project.id)  # 获取最新的项目数据
                # 刷新详情对话框的数据显示
                dialog.refresh_data(upd,
                    self._workflow_service.get_all_stages(),  # 传入最新阶段列表
                    self._log_service.get_project_logs(card.project.id))  # 传入最新日志
            else:
                messagebox.showerror("错误", msg)  # 移动失败弹窗提示

        # 打开项目详情对话框，传入项目数据、阶段列表、日志和移动回调
        result = show_detail_dialog(self, card.project, stages, logs,
                                   on_move=_handle_move)
        if not result:  # 用户关闭了对话框但未执行任何操作
            return

        # 解析对话框返回的操作结果
        action, data = result  # action 为操作类型字符串（"edit" / "delete"），data 为编辑后的表单数据
        if action == "edit":  # 用户在详情窗口中点了"编辑"按钮
            success, msg = self._project_service.update_project(
                card.project.id,  # 要更新的项目 ID
                company_name=data.get("company_name"),  # 公司名称
                system_name=data.get("system_name"),  # 系统名称
                cert_number=data.get("cert_number"),  # 证书编号
                issue_date=data.get("issue_date"),  # 签发日期
                level=data.get("level"),  # 保护等级
                location=data.get("location"),  # 所属地
                deadline=data.get("deadline"),  # 截止日期
                notes=data.get("notes"),  # 备注
                stage_id=data.get("stage_id"),  # 所属阶段
                folder_path=data.get("folder_path"),  # 文件夹路径
            )
            if success:  # 更新成功
                self._refresh_kanban()  # 刷新看板显示最新数据
            else:
                messagebox.showerror("错误", msg)  # 更新失败弹窗提示

        elif action == "delete":  # 用户在详情窗口中点了"删除"按钮
            success, msg = self._project_service.delete_project(card.project.id)  # 执行项目删除
            if success:  # 删除成功
                self._refresh_kanban()  # 刷新看板，移除卡片
            else:
                messagebox.showerror("错误", msg)  # 删除失败弹窗提示

    def _on_card_edit(self, card: ProjectCard):
        """处理卡片"编辑"按钮点击 / 卡片双击事件 - 直接打开项目编辑对话框

        与 _on_card_detail 的区别：
          - _on_card_detail：打开详情窗口（展示完整信息），从详情窗口再进入编辑
          - _on_card_edit：跳过详情窗口，直接打开编辑对话框（快速编辑入口）

        此方法同时作为卡片双击事件的处理函数，提升操作效率。

        Args:
            card: 被双击或被点击"编辑"按钮的项目卡片组件
        """
        stages = self._workflow_service.get_all_stages()  # 获取所有阶段列表（供编辑对话框的下拉选择）
        # 打开编辑对话框，传入当前项目数据作为预填值
        result = show_project_dialog(self, "编辑项目", card.project, stages)
        if result:  # 用户确认修改
            success, msg = self._project_service.update_project(
                card.project.id,  # 要更新的项目唯一 ID
                company_name=result.get("company_name"),  # 公司名称
                system_name=result.get("system_name"),  # 系统名称
                cert_number=result.get("cert_number"),  # 证书编号
                issue_date=result.get("issue_date"),  # 签发日期
                level=result.get("level"),  # 保护等级
                location=result.get("location"),  # 所属地
                deadline=result.get("deadline"),  # 截止日期
                notes=result.get("notes"),  # 备注
                stage_id=result.get("stage_id"),  # 所属阶段
                folder_path=result.get("folder_path"),  # 文件夹路径
            )
            if success:  # 更新成功
                self._refresh_kanban()  # 刷新看板反映最新数据
            else:
                messagebox.showerror("错误", msg)  # 更新失败弹窗提示

    def _on_card_move_stage(self, card: ProjectCard, target_stage_id: str):
        """处理卡片左箭头/右箭头的阶段移动事件

        当用户点击卡片上的 ◀（左箭头）或 ▶（右箭头）时触发。
        左箭头将项目移到上一阶段，右箭头将项目移到下一阶段。

        先执行防御性检查：如果源阶段和目标阶段相同（用户在边界处误点），直接返回。
        成功后直接在 UI 上移动卡片，无需全量刷新看板（性能优化）。

        Args:
            card: 被操作的项目卡片组件
            target_stage_id: 目标阶段的唯一标识符（由 KanbanBoard 计算得出）
        """
        source_stage_id = card.project.stage_id  # 获取卡片当前所在的阶段 ID
        if source_stage_id == target_stage_id:  # 源阶段与目标阶段相同（用户在首/尾误点箭头）
            return  # 无需移动，直接返回

        # 调用项目服务执行阶段移动
        success, msg = self._project_service.move_project(
            card.project.id,  # 要移动的项目 ID
            target_stage_id,  # 目标阶段 ID
        )
        if success:  # 移动成功
            # 直接在 UI 上执行卡片移动（从源列移除，添加到目标列顶部）
            # 这种方式比全量刷新更高效，且不丢失选中状态
            self._kanban.move_card_to_column(card, target_stage_id)
        else:
            messagebox.showerror("错误", msg)  # 移动失败弹窗提示

    def _on_card_copy(self, card: ProjectCard):
        """处理卡片"复制"按钮点击事件 - 创建当前项目的完整副本

        复制逻辑：
          - 读取源项目的所有业务字段（公司名称、系统名称、编号、等级等）
          - 公司名称在原值后追加 " - 副本" 后缀以区分原项目
          - 通过 ProjectService.create_project() 创建新项目实体
          - 新项目与源项目处于同一流程阶段

        Args:
            card: 被点击"复制"按钮的项目卡片组件
        """
        p = card.project  # 获取源项目对象引用（简化后续字段访问）
        copy_name = f"{p.company_name} - \u526f\u672c"  # 公司名称 + " - 副本"（\u526f\u672c = 副本）
        # 调用项目服务创建副本项目，所有字段与源项目保持一致
        success, msg, _ = self._project_service.create_project(
            company_name=copy_name,  # 使用带"副本"后缀的名称
            system_name=p.system_name,  # 系统名称保持不变
            cert_number=p.cert_number,  # 证书编号保持不变
            issue_date=p.issue_date,  # 签发日期保持不变
            level=p.level,  # 保护等级保持不变
            location=p.location,  # 所属地保持不变
            deadline=p.deadline,  # 截止日期保持不变
            notes=p.notes,  # 备注保持不变
            stage_id=p.stage_id,  # 所处阶段保持不变
        )
        if success:  # 复制成功
            self._refresh_kanban()  # 刷新看板，显示新创建的副本卡片
        else:
            messagebox.showerror("错误", msg)  # 复制失败弹窗提示

    def _on_column_resize(self, stage_id: str, new_width: int):
        """处理列宽拖拽完成事件 - 保存调整后的列宽到持久化数据

        当用户拖拽看板列右侧的分隔手柄调整列宽后触发。
        将新宽度通过 WorkflowService 持久化保存，下次启动时自动恢复。

        Args:
            stage_id: 被调整宽度列对应的阶段 ID
            new_width: 新的列宽值（单位：像素）
        """
        self._workflow_service.update_stage_width(stage_id, new_width)  # 保存新列宽到数据服务

    def _create_project_folder(self, project):
        """为新建项目创建完整的文件目录结构

        在程序数据目录下创建以公司名+系统名+日期命名的文件夹，
        并初始化以下子目录：
          - 01-其他归档文件/         # 存放杂项归档材料
          - 00-{公司}-{系统}-报告打印/  # 存放打印版报告
          - 13-{公司}-{系统}-渗透测试报告/  # 存放渗透测试相关文件

        文件夹命名规则：
          格式：{序号}-{公司简称}-{系统简称}-{年月日}
          序号：当前项目总数（自动递增）
          日期：创建当日（格式 YYMMDD）

        创建成功后，将文件夹路径保存到 project 对象并持久化到 JSON 文件。

        Args:
            project: 新创建的项目实体对象（包含公司名称、系统名称等信息）
        """
        import os  # 操作系统接口，用于创建目录和处理路径
        from datetime import date  # 日期类，用于生成日期字符串

        try:
            base = Config.get_data_dir()  # 获取程序数据根目录路径
            count = len(self._project_service.get_all_projects())  # 当前项目总数（含刚创建的）
            date_str = date.today().strftime("%y%m%d")  # 当日日期，如 "260603"（年月日格式）
            # 清理公司名称中的非法字符（路径分隔符）
            cname = (project.company_name or "未命名").replace("/", "_").replace("\\", "_")
            # 清理系统名称中的非法字符
            sname = (project.system_name or "").replace("/", "_").replace("\\", "_")
            # 组装文件夹名称：序号-公司名-系统名-日期
            folder_name = f"{count:03d}-{cname}-{sname}-{date_str}"
            root = os.path.join(base, folder_name)  # 拼接完整文件夹路径
            os.makedirs(root, exist_ok=True)  # 递归创建文件夹目录（exist_ok=True 避免重复创建报错）

            # 定义需要创建的子目录列表
            subdirs = [
                "01-其他归档文件",
                f"00-{cname}-{sname}-报告打印",
            ]
            for d in subdirs:
                os.makedirs(os.path.join(root, d), exist_ok=True)

            # 生成保密承诺书模板（替换年份和公司名称）
            self._generate_nda_template(root, cname, project.company_name or "未命名")

            # 持久化文件夹路径
            project.folder_path = root  # 设置项目对象的文件夹路径属性
            self._data_service.update_project(project.id, {"folder_path": root})  # 持久化到 JSON
        except OSError:
            pass

    def _generate_nda_template(self, root, cname_clean, company_name):
        """生成保密承诺书模板，替换公司名称和日期。

        Args:
            root: 项目文件夹根路径
            cname_clean: 清理后的公司名称（用于文件名）
            company_name: 原始公司名称（用于文档内容替换）
        """
        import os, shutil
        try:
            template_path = os.path.join(Config.get_data_dir(), "templates", "02-保密承诺书模板.docx")
            if not os.path.exists(template_path):
                return
            dest_name = f"02-{cname_clean}-{company_name}-保密承诺书.docx" if company_name else f"02-{cname_clean}-保密承诺书.docx"
            dest_path = os.path.join(root, dest_name)
            if os.path.exists(dest_path):
                return
            shutil.copy2(template_path, dest_path)
            import docx
            doc = docx.Document(dest_path)
            # 替换所有 XX公司/XXX公司 → 实际公司名
            for p in doc.paragraphs:
                for run in p.runs:
                    if "公司" in run.text and ("XX" in run.text or "XXX" in run.text):
                        run.text = run.text.replace("XXX公司", company_name).replace("XX公司", company_name)
            from datetime import datetime
            create_date = datetime.now().strftime("%Y年%m月%d日")
            # 替换所有 XX年XX月XX日 → 创建日期
            for p in doc.paragraphs:
                for run in p.runs:
                    run.text = run.text.replace("XX年XX月XX日", create_date)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.text = run.text.replace("XX年XX月XX日", create_date)
            doc.save(dest_path)
        except Exception:
            pass

    # ==================================================================================
    # 窗口事件处理方法
    # ==================================================================================

    def _on_window_resize(self, event):
        """处理窗口大小变化事件（预留扩展）

        Tkinter 的 pack 布局管理器会自动处理组件的大小调整，
        此方法作为预留扩展点，可在此添加额外的自适应逻辑，
        例如：当窗口宽度低于阈值时折叠工具栏。

        Args:
            event: Tkinter 的 Configure 事件对象（包含新窗口尺寸等信息）
        """
        pass  # 当前无需额外处理，pack 布局自动完成自适应

    def _on_close(self):
        """窗口关闭前保存数据并询问是否同步到 WebDAV。"""
        self._data_service.save()
        # 询问是否备份到 WebDAV
        from utils.webdav_config import WebDAVConfig
        cfg = WebDAVConfig.load()
        if cfg.url:
            if messagebox.askyesno("数据同步", "是否将当前数据同步到 WebDAV 服务器？"):
                try:
                    from services.backup_service import BackupService
                    svc = BackupService(cfg)
                    ok, msg = svc.backup(Config.get_data_file_path())
                    if ok:
                        messagebox.showinfo("同步成功", f"数据已备份: {msg}")
                    else:
                        messagebox.showwarning("同步失败", msg)
                except Exception as e:
                    messagebox.showwarning("同步错误", str(e))
        self.destroy()

    def _check_restore_on_startup(self):
        """启动时检查 WebDAV 备份，提示是否恢复数据。"""
        from utils.webdav_config import WebDAVConfig
        cfg = WebDAVConfig.load()
        if not cfg.url:
            return
        try:
            from services.backup_service import BackupService
            svc = BackupService(cfg)
            ok, msg, files = svc.list_backups()
            if not ok or not files:
                return
            files.sort(key=lambda x: x["name"], reverse=True)
            # 构建备份列表对话框
            import tkinter as tk
            dlg = tk.Toplevel(self)
            dlg.title("数据恢复 - 检测到云端备份")
            dlg.geometry("550x400")
            dlg.configure(bg="#ffffff")
            dlg.grab_set()
            tk.Label(dlg, text="检测到以下云端备份，是否恢复？", bg="#ffffff",
                     font=("Microsoft YaHei", 12, "bold"), fg="#2c3e50",
                     ).pack(pady=(15, 10))
            frame = tk.Frame(dlg, bg="#ffffff")
            frame.pack(fill=tk.BOTH, expand=True, padx=15, pady=5)
            lb = tk.Listbox(frame, font=("Microsoft YaHei", 10), selectmode="single")
            lb.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
            sb = tk.Scrollbar(frame, orient=tk.VERTICAL, command=lb.yview)
            sb.pack(side=tk.RIGHT, fill=tk.Y)
            lb.configure(yscrollcommand=sb.set)
            for f in files[:20]:
                lb.insert(tk.END, f"{f['name']}  ({f.get('modified','?')})")
            result = {"selected": False}
            def _restore():
                sel = lb.curselection()
                if not sel:
                    messagebox.showwarning("提示", "请选择要恢复的备份", parent=dlg)
                    return
                idx = sel[0]
                if messagebox.askyesno("确认恢复", f"确定要恢复「{files[idx]['name']}」吗？\n当前数据将被覆盖！", parent=dlg):
                    ok2, msg2, body = svc.restore(files[idx]["path"])
                    if ok2:
                        import json
                        try:
                            json.loads(body.decode("utf-8"))
                        except Exception:
                            messagebox.showerror("错误", "备份文件格式错误", parent=dlg)
                            return
                        with open(Config.get_data_file_path(), "wb") as wf:
                            wf.write(body)
                        self._data_service.reload()
                        self._refresh_kanban()
                        result["selected"] = True
                        messagebox.showinfo("成功", "数据已恢复", parent=dlg)
                        dlg.destroy()
                    else:
                        messagebox.showerror("恢复失败", msg2, parent=dlg)
            btn_frame = tk.Frame(dlg, bg="#f0f2f5")
            btn_frame.pack(fill=tk.X, side=tk.BOTTOM)
            tk.Frame(btn_frame, bg="#d0d5dd", height=1).pack(fill=tk.X)
            inner = tk.Frame(btn_frame, bg="#f0f2f5")
            inner.pack(fill=tk.X, padx=16, pady=8)
            tk.Button(inner, text="跳过", command=dlg.destroy,
                bg="#ffffff", fg="#2c3e50", cursor="hand2",
                font=("Microsoft YaHei", 10), relief="flat", padx=18, pady=5,
                highlightbackground="#d0d5dd", highlightthickness=1,
                ).pack(side=tk.RIGHT, padx=(10, 0))
            tk.Button(inner, text="恢复选中", command=_restore,
                bg="#3498db", fg="white", cursor="hand2",
                font=("Microsoft YaHei", 10, "bold"), relief="flat", padx=18, pady=5,
                ).pack(side=tk.RIGHT)
            self.wait_window(dlg)
        except Exception:
            pass

    # ==================================================================================
    # 内部辅助方法
    # ==================================================================================

    def _refresh_kanban(self):
        """刷新看板显示 - 从服务层重新加载所有数据并重建 UI

        这是数据变更后的统一 UI 刷新入口，所有导致看板内容变化的操作
        （新增、编辑、删除、移动项目，编辑流程等）最终都通过此方法刷新显示。

        执行流程：
          1. 从 WorkflowService 获取最新阶段列表
          2. 从 ProjectService 获取最新项目列表
          3. 调用 KanbanBoard.load_stages() 销毁旧组件、创建新列和卡片
        """
        stages = self._workflow_service.get_all_stages()  # 从流程服务获取最新阶段列表（含颜色、宽度）
        projects = self._project_service.get_all_projects()  # 从项目服务获取最新项目列表（含所有字段）
        self._kanban.load_stages(stages, projects)  # 看板重建所有列和卡片，完成渲染
